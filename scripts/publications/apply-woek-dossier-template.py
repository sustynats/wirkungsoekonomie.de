#!/usr/bin/env python3
"""Apply the WÖk dossier/concept Word template to a source DOCX.

The script keeps the source document order and content, but wraps it in the
public WÖk publication shell: cover, metadata, leitformel, publication note,
generated table of contents and template styles.
"""

from __future__ import annotations

import argparse
import io
import re
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph


DEFAULT_TEMPLATE = Path(
    "/Users/hagen/Documents/Rechner-Cleanup/Sortiert_2026-05-25/"
    "01_WOeK/01_Kerndokumente-Konzepte/Pakete-Ordner/Downloads/"
    "WOeK_Dossier_Template_Paket/WOeK_Dossier_Konzept_Template.dotx"
)

CONTENT_TYPE_TEMPLATE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
)
CONTENT_TYPE_DOCUMENT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)


def compatible_docx_path(path: Path) -> tempfile.NamedTemporaryFile | None:
    """Return a temporary DOCX for DOTX input, otherwise None."""
    if path.suffix.lower() != ".dotx":
        return None

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        tmp.name, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    CONTENT_TYPE_TEMPLATE.encode("utf-8"),
                    CONTENT_TYPE_DOCUMENT.encode("utf-8"),
                )
            target.writestr(item, data)
    return tmp


def open_document(path: Path) -> tuple[Document, tempfile.NamedTemporaryFile | None]:
    tmp = compatible_docx_path(path)
    return Document(tmp.name if tmp else str(path)), tmp


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def text_of_element(element) -> str:
    return "".join(t.text or "" for t in element.xpath(".//w:t"))


def replace_text_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    text = "".join(run.text for run in paragraph.runs)
    new_text = text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    new_text = re.sub(r"\{\{[^{}]+\}\}", "", new_text)
    if new_text == text:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_placeholders(document: Document, replacements: dict[str, str]) -> None:
    stories = [document]
    for section in document.sections:
        stories.extend([section.header, section.footer])
    for story in stories:
        for paragraph in story.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for table in story.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)


def remove_elements_after_toc(document: Document) -> None:
    body = document.element.body
    keep_until = None
    for idx, child in enumerate(list(body)):
        if child.tag == qn("w:p") and text_of_element(child).strip() == "Inhaltsverzeichnis":
            keep_until = idx
            break
    if keep_until is None:
        return
    for child in list(body)[keep_until + 1 :]:
        body.remove(child)


def add_generated_toc(target: Document, headings: list[tuple[int, str]]) -> None:
    if not headings:
        return
    note = target.add_paragraph(
        "Automatisch erzeugte Orientierung aus den Überschriften der Online- und "
        "Downloadfassung. Seitenzahlen werden beim Öffnen in Word aktualisiert."
    )
    note.style = "WÖk Kleindruck" if "WÖk Kleindruck" in target.styles else "Normal"
    for level, title in headings:
        if title == "Inhaltsübersicht":
            continue
        paragraph = target.add_paragraph()
        paragraph.style = "Normal"
        if level > 1:
            paragraph.paragraph_format.left_indent = 457200 * (level - 1)
        paragraph.add_run(title)
    target.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def source_start_index(blocks: list[Paragraph | Table], start_heading: str) -> int:
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.text.strip() == start_heading:
            return idx
    return 0


def collect_headings(blocks: list[Paragraph | Table], start: int) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for block in blocks[start:]:
        if not isinstance(block, Paragraph):
            continue
        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style else ""
        match = re.match(r"Heading (\d+)", style)
        if match:
            headings.append((int(match.group(1)), text))
    return headings


def remap_relationships(element, source: Document, target: Document) -> None:
    rel_attr_names = [
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
    ]
    for node in element.iter():
        for attr_name in rel_attr_names:
            old_rid = node.attrib.get(attr_name)
            if not old_rid or old_rid not in source.part.rels:
                continue
            rel = source.part.rels[old_rid]
            if rel.reltype == RT.IMAGE:
                new_rid, _ = target.part.get_or_add_image(io.BytesIO(rel.target_part.blob))
                node.attrib[attr_name] = new_rid
            elif rel.reltype == RT.HYPERLINK:
                new_rid = target.part.relate_to(rel.target_ref, RT.HYPERLINK, True)
                node.attrib[attr_name] = new_rid


def append_source_content(source: Document, target: Document, start_heading: str) -> None:
    blocks = list(iter_block_items(source))
    start = source_start_index(blocks, start_heading)
    add_generated_toc(target, collect_headings(blocks, start))
    body = target.element.body
    sect_pr = body.sectPr
    for block in blocks[start:]:
        if isinstance(block, Paragraph) and block.text.strip() == "Inhaltsübersicht":
            continue
        cloned = deepcopy(block._element)
        remap_relationships(cloned, source, target)
        if sect_pr is not None:
            body.insert(body.index(sect_pr), cloned)
        else:
            body.append(cloned)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--document-type", default="Diskussionspapier")
    parser.add_argument("--title", required=True)
    parser.add_argument("--subtitle", required=True)
    parser.add_argument("--version", default="v1.0")
    parser.add_argument("--fassung", default="Diskussionsfassung")
    parser.add_argument("--stand", default="Juni 2026")
    parser.add_argument("--kurztitel", default="")
    parser.add_argument("--author", default="Natalie Weber")
    parser.add_argument("--reference", default="Wirkungsökonomie")
    parser.add_argument("--geltung", default="Öffentliche Konzept- und Dossierfassung")
    parser.add_argument(
        "--publication-note",
        default=(
            "Dieses Dokument ist als öffentliche Diskussionsfassung gesetzt. "
            "Die fachliche Reihenfolge und der Inhalt bleiben bei der "
            "Standardisierung erhalten; angepasst werden Layout, Formatvorlagen, "
            "Titelblatt, Kopf-/Fußzeilen, Tabellenoptik und typografische Konsistenz."
        ),
    )
    parser.add_argument("--start-heading", default="Dokumentenstatus und Zweck")
    args = parser.parse_args()

    if not args.source.exists():
        raise SystemExit(f"Source DOCX not found: {args.source}")
    if not args.template.exists():
        raise SystemExit(f"WÖk template not found: {args.template}")

    template_doc, tmp_template = open_document(args.template)
    source_doc = Document(str(args.source))
    replacements = {
        "{{DOKUMENTTYP}}": args.document_type,
        "{{TITEL DES DOKUMENTS}}": args.title,
        "{{UNTERTITEL / EINORDNUNG}}": args.subtitle,
        "{{VERSION}}": args.version,
        "{{FASSUNG}}": args.fassung,
        "{{MONAT JAHR}}": args.stand,
        "{{KURZTITEL}}": args.kurztitel or args.title,
        "{{AUTORIN}}": args.author,
        "{{REFERENZ / PROJEKT}}": args.reference,
        "{{Öffentliche Konzept- und Dossierfassung / Interne Arbeitsfassung}}": args.geltung,
        "{{STAND}}": args.stand,
        "{{Diskussionsfassung / freigegeben / Entwurf}}": args.fassung,
        "{{Arbeitspapier / Konzeptpapier / Dossier}}": args.document_type,
        "Template für Arbeitspapiere, Konzepte und Dossiers": args.geltung,
        "Nicht verwenden für Präsentationen, Buch-/Langformate, Manifest, Minifest, Parteiprogramme oder Presseartikel.": (
            "Hinweis: Dieses Papier ist kein amtlicher Index und keine validierte Statistik. "
            "Es ist ein Konzeptvorschlag, der bestehende Datenquellen, SDG-Indikatoren "
            "und wirkungsökonomische Logik zu einem prüfbaren Modell verbindet."
        ),
        "Dieses Dokument ist als {{Arbeitspapier / Konzeptpapier / Dossier}} gesetzt. Die bestehende Reihenfolge und der fachliche Inhalt bleiben bei der Standardisierung unverändert; angepasst werden nur Layout, Formatvorlagen, Titelblatt, Kopf-/Fußzeilen, Tabellenoptik und typografische Konsistenz.": args.publication_note,
    }

    replace_placeholders(template_doc, replacements)
    remove_elements_after_toc(template_doc)
    append_source_content(source_doc, template_doc, args.start_heading)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    template_doc.save(args.output)
    if tmp_template:
        Path(tmp_template.name).unlink(missing_ok=True)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
