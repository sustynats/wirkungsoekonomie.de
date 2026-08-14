#!/usr/bin/env python3
"""Build the public v1.2 edition of the WÖk Value Pricing dossier.

The supplied Word file remains a non-public source asset.  This builder makes
the publication reproducible: it removes the source cover and duplicate manual
table of contents, applies the established WÖk dossier template, creates a
clean PDF, and derives the readable online edition from the same source.
"""
from __future__ import annotations

import html
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "source-assets/originals/Dossier_Wirkungsbasiertes_Value_Pricing_v1.2.docx"
TEMPLATE = ROOT / (
    "Sortiert_2026-05-25/01_WOeK/01_Kerndokumente-Konzepte/Pakete-Ordner/Downloads/"
    "WOeK_Dossier_Template_Paket/WOeK_Dossier_Konzept_Template.dotx"
)
ONLINE = ROOT / "content/documents/online/wirkungsbasiertes-value-pricing-v1.2.inc"
PDF = ROOT / "public/downloads/originals/Wirkungsbasiertes_Value_Pricing_v1.2.pdf"
WORKDIR = ROOT / "tmp/publications/wirkungsbasiertes-value-pricing-v1.2"
PREPARED = WORKDIR / "wirkungsbasiertes-value-pricing-v1.2-prepared.docx"
PUBLIC_DOCX = WORKDIR / "wirkungsbasiertes-value-pricing-v1.2-public.docx"
TITLE = "Wirkungsbasiertes Value Pricing"
SUBTITLE = "Vom Stundenhonorar zur Vergütung nach erzeugtem Kundenwert und positiver Netto-Wirkung"
EDITION = "Version 1.2 · Erstveröffentlichung 4. Mai 2026 · aktualisierte Fassung August 2026"
AI_MARKER = re.compile(
    r"\b(?:chatgpt|openai|claude|anthropic|gemini|copilot|codex|code[xs]|system\s*prompt|user\s*prompt|ki-?anweisung)\b"
    r"|(?:utm_(?:source|medium|campaign)=)(?:chatgpt|openai|claude|anthropic|gemini|copilot)",
    re.IGNORECASE,
)


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def is_heading(block: Paragraph, text: str) -> bool:
    return isinstance(block, Paragraph) and block.text.strip() == text and block.style.name.startswith("Heading")


def delete_block(block: Paragraph | Table) -> None:
    element = block._element
    element.getparent().remove(element)


def sanitize_url(value: str) -> str:
    """Remove assistant-tracking query values while preserving real citations."""
    try:
        parsed = urlsplit(value)
    except ValueError:
        return value
    if not parsed.scheme or not parsed.netloc:
        return value
    clean_pairs = [
        (key, item)
        for key, item in parse_qsl(parsed.query, keep_blank_values=True)
        if not (
            key.lower() in {"utm_source", "utm_medium", "utm_campaign"}
            and item.lower().removesuffix(".com") in {"chatgpt", "openai", "claude", "anthropic", "gemini", "copilot"}
        )
    ]
    return urlunsplit((parsed.scheme, parsed.netloc, parsed.path, urlencode(clean_pairs, doseq=True), parsed.fragment))


def clean_relationship_targets(docx_path: Path) -> None:
    """Sanitize external link targets and remove comments/custom properties."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                name = item.filename
                if name in {"docProps/custom.xml", "word/comments.xml", "word/commentsExtended.xml", "word/people.xml"}:
                    continue
                data = source.read(name)
                if name.endswith(".rels"):
                    root = ET.fromstring(data)
                    changed = False
                    for relationship in root:
                        target_value = relationship.attrib.get("Target", "")
                        clean = sanitize_url(target_value)
                        if clean != target_value:
                            relationship.set("Target", clean)
                            changed = True
                    if changed:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                target.writestr(item, data)
        temporary.replace(docx_path)
    finally:
        temporary.unlink(missing_ok=True)


def clean_source() -> Document:
    source = Document(SOURCE)
    # The public template supplies a generated contents overview.  The hand-set
    # source directory would otherwise duplicate all navigation in the PDF.
    blocks = list(iter_blocks(source))
    content_start = next((index for index, block in enumerate(blocks) if is_heading(block, "Inhalt")), None)
    first_chapter = next(
        (index for index, block in enumerate(blocks)
         if isinstance(block, Paragraph) and block.style.name.startswith("Heading") and re.match(r"1\s+Zweck", block.text.strip())),
        None,
    )
    if content_start is not None and first_chapter is not None and first_chapter > content_start:
        for block in blocks[content_start:first_chapter]:
            delete_block(block)

    # This is an internal publication hand-off, not part of the public
    # argument.  It must never appear in the online edition or the PDF.
    blocks = list(iter_blocks(source))
    editorial_start = next(
        (
            index
            for index, block in enumerate(blocks)
            if isinstance(block, Paragraph)
            and block.text.strip() == "15.5 Website- und Publikationsarchitektur"
        ),
        None,
    )
    editorial_end = next(
        (
            index
            for index, block in enumerate(blocks)
            if editorial_start is not None
            and index > editorial_start
            and isinstance(block, Paragraph)
            and block.text.strip() == "15.6 Forschungsfragen"
        ),
        None,
    )
    if editorial_start is None or editorial_end is None:
        raise RuntimeError("Redaktionellen Abschnitt 15.5 im Dossier nicht eindeutig gefunden.")
    for block in blocks[editorial_start:editorial_end]:
        delete_block(block)
    return source


def public_master_from_source(source: Document) -> None:
    """Keep the supplied WÖk master cover when the optional DOTX is unavailable."""
    for table in source.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Version 1.2" not in paragraph.text:
                        continue
                    for run in paragraph.runs:
                        if "Version 1.2" in run.text:
                            run.text = re.sub(
                                r"Version\s*1\.2\s*[·•-]?\s*August\s*2026",
                                "Version 1.2 · Erstveröffentlichung 4. Mai 2026 · aktualisierte Fassung August 2026",
                                run.text,
                            )


def style_level(paragraph: Paragraph) -> int | None:
    match = re.match(r"Heading\s+(\d+)", paragraph.style.name or "")
    return int(match.group(1)) if match else None


def slug(value: str, used: dict[str, int]) -> str:
    token = value.lower().replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    token = re.sub(r"[^a-z0-9]+", "-", token).strip("-") or "abschnitt"
    used[token] = used.get(token, 0) + 1
    return token if used[token] == 1 else f"{token}-{used[token]}"


def render_runs(paragraph: Paragraph) -> str:
    bits: list[str] = []
    for run in paragraph.runs:
        text = html.escape(run.text or "").replace("\n", "<br>")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        bits.append(text)
    raw = "".join(bits) or html.escape(paragraph.text or "")
    return re.sub(r"https?://[^\s<]+", lambda match: f'<a href="{html.escape(sanitize_url(match.group(0).rstrip(".,;:)")), quote=True)}">{html.escape(match.group(0).rstrip(".,;:)"))}</a>{html.escape(match.group(0)[len(match.group(0).rstrip(".,;:)")):])}', raw)


def render_table(table: Table) -> str:
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            value = "<br>".join(render_runs(paragraph) for paragraph in cell.paragraphs if paragraph.text.strip()).strip() or "&nbsp;"
            tag = "th" if row_index == 0 else "td"
            scope = ' scope="col"' if row_index == 0 else ""
            cells.append(f"<{tag}{scope}>{value}</{tag}>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    if not rows:
        return ""
    return f'<div class="table-scroll"><table class="data-table"><thead>{rows[0]}</thead><tbody>{"".join(rows[1:])}</tbody></table></div>'


def build_online(source: Document) -> str:
    used: dict[str, int] = {}
    output = [
        f"<p><b>Autorin:</b> Natalie Weber · <b>{html.escape(EDITION)}</b> · <b>Status:</b> Fachliches Dossier und Arbeitsmodell</p>",
        "<blockquote><p><b>Ausgabehinweis:</b> Diese Online-Lesefassung erschließt den vollständigen Text. Die gestaltete PDF-Fassung enthält zusätzlich die gesetzten Abbildungen, Formeln und Tabellen im WÖk-Dossier-Design.</p></blockquote>",
    ]
    start = False
    for block in iter_blocks(source):
        if isinstance(block, Paragraph) and block.text.strip() == "Publikationshinweis":
            start = True
        if not start:
            continue
        if isinstance(block, Table):
            output.append(render_table(block))
            continue
        text = block.text.strip()
        if not text:
            continue
        level = style_level(block)
        if level:
            tag = min(max(level + 1, 2), 4)
            output.append(f'<h{tag} id="{slug(text, used)}">{render_runs(block)}</h{tag}>')
            continue
        style_name = block.style.name or ""
        if style_name.startswith("List"):
            output.append(f"<ul><li>{render_runs(block)}</li></ul>")
        else:
            output.append(f"<p>{render_runs(block)}</p>")
    output.extend([
        '<h2 id="quellenarchiv">Quellenarchiv</h2>',
        '<p>Die zentralen externen Referenzen dieser Veröffentlichung sind im Quellenarchiv mit ihrem Veröffentlichungsbezug hinterlegt.</p>',
        '<ul>'
        '<li><a href="/quellenarchiv/wok-q-1091/">Hinterhuber (2004): Towards value-based pricing</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1092/">Töytäri und Rajala (2015): Value-based selling</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1093/">Töytäri, Keränen und Rajala (2017): Barriers to implementing value-based pricing</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1094/">Raja et al. (2020): Learning to discover value</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1095/">Keränen et al. (2023): Gain-sharing in performance-based contracting</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1096/">Selviaridis und Wynstra (2015): Performance-based contracting</a></li>'
        '<li><a href="/quellenarchiv/wok-q-1097/">Ng, Maull und Yip (2009): Outcome-based contracts</a></li>'
        '</ul>',
    ])
    return "\n".join(output) + "\n"


def ensure_no_residue(paths: list[Path]) -> None:
    findings: list[str] = []
    for path in paths:
        if not path.exists():
            continue
        if path.suffix == ".docx":
            with zipfile.ZipFile(path) as archive:
                text = "\n".join(
                    archive.read(name).decode("utf-8", errors="ignore")
                    for name in archive.namelist()
                    if name.endswith((".xml", ".rels"))
                )
        else:
            text = path.read_text(encoding="utf-8")
        if AI_MARKER.search(text):
            findings.append(str(path.relative_to(ROOT)))
    if findings:
        raise RuntimeError("Nicht veröffentlichungsfähige KI-/Produktionsreste: " + ", ".join(findings))


def export_pdf(docx_path: Path) -> None:
    soffice = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if not soffice.exists():
        raise RuntimeError("LibreOffice für den PDF-Export nicht gefunden.")
    PDF.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="woek-value-pricing-pdf-") as temp:
        out = Path(temp)
        profile = out / "profile"
        profile.mkdir()
        result = subprocess.run(
            [str(soffice), "--headless", f"-env:UserInstallation=file://{profile}", "--convert-to", "pdf", "--outdir", str(out), str(docx_path)],
            check=False,
            capture_output=True,
            text=True,
        )
        exported = out / f"{docx_path.stem}.pdf"
        if result.returncode != 0 or not exported.exists():
            raise RuntimeError(f"LibreOffice-PDF-Export fehlgeschlagen: {result.stderr.strip()}")
        shutil.copyfile(exported, PDF)


def main() -> None:
    if not SOURCE.exists():
        raise RuntimeError("Dossierquelle fehlt.")
    WORKDIR.mkdir(parents=True, exist_ok=True)
    source = clean_source()
    source.save(PREPARED)
    clean_relationship_targets(PREPARED)
    ONLINE.parent.mkdir(parents=True, exist_ok=True)
    ONLINE.write_text(build_online(Document(PREPARED)), encoding="utf-8")
    if TEMPLATE.exists():
        subprocess.run(
            [
                sys.executable,
                str(ROOT / "scripts/publications/apply-woek-dossier-template.py"),
                str(PREPARED),
                str(PUBLIC_DOCX),
                "--template", str(TEMPLATE),
                "--document-type", "Dossier",
                "--title", TITLE,
                "--subtitle", SUBTITLE,
                "--version", "v1.2",
                "--fassung", "Fachliches Dossier und Arbeitsmodell",
                "--stand", "Erstveröffentlichung 4. Mai 2026 · aktualisierte Fassung August 2026",
                "--kurztitel", TITLE,
                "--author", "Natalie Weber",
                "--reference", "Wirkungsökonomie",
                "--geltung", "Öffentliches Dossier und fachliches Detailkonzept",
                "--publication-note", "Diese öffentliche Fassung nutzt das WÖk-Dossier-Design. Sie dokumentiert ein fachliches Arbeitsmodell; es ist kein allgemein anerkannter Norm-, Prüf- oder Vertragsstandard und keine Finanz-, Rechts- oder Preisberatung.",
                "--start-heading", "Publikationshinweis",
            ],
            check=True,
        )
    else:
        # The delivered v1.2 source is already the approved WÖk Dossier master.
        # Keep its page geometry, figures and typography rather than creating a
        # substitute layout when the optional, untracked DOTX package is absent.
        public_master = Document(PREPARED)
        public_master_from_source(public_master)
        public_master.save(PUBLIC_DOCX)
    clean_relationship_targets(PUBLIC_DOCX)
    public = Document(PUBLIC_DOCX)
    public.core_properties.title = TITLE
    public.core_properties.subject = "Dossier und fachliches Detailkonzept der Wirkungsökonomie"
    public.core_properties.author = "Natalie Weber"
    public.core_properties.keywords = "Wirkungsökonomie, Value-based Pricing, Kundenwert, Netto-Wirkung"
    public.core_properties.comments = "Öffentliche Fassung v1.2"
    public.core_properties.last_modified_by = "Wirkungsökonomie"
    public.save(PUBLIC_DOCX)
    clean_relationship_targets(PUBLIC_DOCX)
    ensure_no_residue([PREPARED, PUBLIC_DOCX, ONLINE])
    export_pdf(PUBLIC_DOCX)
    print(f"wrote {ONLINE.relative_to(ROOT)}")
    print(f"wrote {PUBLIC_DOCX.relative_to(ROOT)}")
    print(f"wrote {PDF.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
