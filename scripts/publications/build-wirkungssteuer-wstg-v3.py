#!/usr/bin/env python3
"""Build the complete public online edition of Wirkungssteuer/WStG 3.0.

The supplied, visually reviewed DOCX remains the non-public publication master.
The branded PDF is committed separately. This builder derives an accessible,
link-preserving HTML include so the publication remains fully readable online.
"""
from __future__ import annotations

import html
import re
from pathlib import Path
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "source-assets/originals/Wirkungssteuer_WStG_3.0_Gesamtneufassung_2026.docx"
ONLINE = ROOT / "content/documents/online/wirkungssteuer-wstg-3-0.inc"
TITLE = "Wirkungssteuer 3.0 und Wirkungssteuergesetz (WStG) 3.0"
EDITION = "Gesamtneufassung 2026 · Version 3.0 · Stand 17. August 2026"
AI_MARKER = re.compile(
    r"\b(?:chatgpt|openai|claude|anthropic|gemini|copilot|codex|system\s*prompt|user\s*prompt|ki-?anweisung)\b"
    r"|(?:utm_(?:source|medium|campaign)=)(?:chatgpt|openai|claude|anthropic|gemini|copilot)",
    re.IGNORECASE,
)


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def sanitize_url(value: str) -> str:
    try:
        parsed = urlsplit(value)
    except ValueError:
        return value
    if not parsed.scheme or not parsed.netloc:
        return value
    pairs = [
        (key, item)
        for key, item in parse_qsl(parsed.query, keep_blank_values=True)
        if not (
            key.lower() in {"utm_source", "utm_medium", "utm_campaign"}
            and item.lower().removesuffix(".com")
            in {"chatgpt", "openai", "claude", "anthropic", "gemini", "copilot"}
        )
    ]
    return urlunsplit((parsed.scheme, parsed.netloc, parsed.path, urlencode(pairs, doseq=True), parsed.fragment))


def run_html(run_element) -> str:
    text_parts: list[str] = []
    for node in run_element.iter():
        if node.tag == qn("w:t"):
            text_parts.append(node.text or "")
        elif node.tag == qn("w:tab"):
            text_parts.append(" ")
        elif node.tag == qn("w:br"):
            text_parts.append("<br>")
    value = html.escape("".join(text_parts)).replace("&lt;br&gt;", "<br>")
    if not value:
        return ""
    properties = run_element.find(qn("w:rPr"))
    if properties is not None:
        if properties.find(qn("w:b")) is not None:
            value = f"<strong>{value}</strong>"
        if properties.find(qn("w:i")) is not None:
            value = f"<em>{value}</em>"
    return value


def inline_html(paragraph: Paragraph) -> str:
    output: list[str] = []
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            output.append(run_html(child))
        elif child.tag == qn("w:hyperlink"):
            label = "".join(run_html(run) for run in child if run.tag == qn("w:r"))
            rel_id = child.get(qn("r:id"))
            target = paragraph.part.rels[rel_id].target_ref if rel_id and rel_id in paragraph.part.rels else ""
            if target:
                clean = sanitize_url(target)
                external = clean.startswith(("http://", "https://"))
                attrs = ' target="_blank" rel="noopener noreferrer"' if external else ""
                output.append(f'<a class="text-link" href="{html.escape(clean, quote=True)}"{attrs}>{label}</a>')
            else:
                output.append(label)
    value = "".join(output).strip()
    return value or html.escape(paragraph.text.strip())


def heading_level(paragraph: Paragraph) -> int | None:
    match = re.match(r"Heading\s+(\d+)", paragraph.style.name or "")
    return int(match.group(1)) if match else None


def slug(value: str, used: dict[str, int]) -> str:
    token = value.lower().replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    token = re.sub(r"[^a-z0-9]+", "-", token).strip("-") or "abschnitt"
    used[token] = used.get(token, 0) + 1
    return token if used[token] == 1 else f"{token}-{used[token]}"


def render_table(table: Table) -> str:
    rows: list[str] = []
    for row_index, row in enumerate(table.rows):
        cells: list[str] = []
        for cell in row.cells:
            value = "<br>".join(inline_html(p) for p in cell.paragraphs if p.text.strip()) or "&nbsp;"
            tag = "th" if row_index == 0 else "td"
            scope = ' scope="col"' if row_index == 0 else ""
            cells.append(f"<{tag}{scope}>{value}</{tag}>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    if not rows:
        return ""
    return f'<div class="table-scroll"><table class="data-table"><thead>{rows[0]}</thead><tbody>{"".join(rows[1:])}</tbody></table></div>'


def build_online(document: Document) -> str:
    used: dict[str, int] = {}
    output = [
        f"<p><strong>Autorin:</strong> Natalie Weber · <strong>{html.escape(EDITION)}</strong> · <strong>Status:</strong> Referenz- und Diskussionsentwurf, kein geltendes Recht</p>",
        "<blockquote><p><strong>Ausgabehinweis:</strong> Diese vollständige Online-Lesefassung erschließt alle 125 kommentierten Paragrafen, Tabellen und Quellen. Die gestaltete PDF-Fassung bewahrt zusätzlich den verbindlichen Satz und die Seitenzählung.</p></blockquote>",
        "<blockquote><p><strong>Leitende Korrektur:</strong> Wirkungssteuer 3.0 ist keine universelle Supersteuer. Sie ist eine modulare Rückkopplungsarchitektur über bestehende Steuerarten. Scorecards liefern Evidenz; Steuersätze und Rechtsfolgen bestimmt weiterhin der demokratisch legitimierte Gesetzgeber.</p></blockquote>",
    ]
    started = False
    list_items: list[str] = []

    def flush_list() -> None:
        nonlocal list_items
        if list_items:
            output.append(f"<ul>{''.join(f'<li>{item}</li>' for item in list_items)}</ul>")
            list_items = []

    for block in iter_blocks(document):
        if isinstance(block, Table):
            if not started:
                continue
            flush_list()
            output.append(render_table(block))
            continue
        text = block.text.strip()
        level = heading_level(block)
        if not started:
            if level == 1 and text == "Inhaltsübersicht":
                started = True
            else:
                continue
        if not text:
            flush_list()
            continue
        if AI_MARKER.search(text):
            raise RuntimeError(f"KI-/Redaktionsrest in öffentlicher Quelle gefunden: {text[:160]}")
        if level:
            flush_list()
            tag = min(max(level + 1, 2), 4)
            output.append(f'<h{tag} id="{slug(text, used)}">{inline_html(block)}</h{tag}>')
            continue
        rendered = inline_html(block)
        if text.startswith("•"):
            # Word may keep the bullet in its own bold run. Remove the glyph
            # together with that optional wrapper so the semantic <ul> does
            # not render a second visible bullet.
            rendered = re.sub(r"^(?:<strong>)?•\s*(?:</strong>)?", "", rendered).strip()
            list_items.append(rendered)
            continue
        flush_list()
        style = block.style.name or ""
        if style == "Lead":
            output.append(f'<p class="lead">{rendered}</p>')
        elif style == "Commentary":
            output.append(f'<div class="callout"><p>{rendered}</p></div>')
        elif style in {"Rationale", "LegalNote"}:
            output.append(f'<div class="callout warning"><p>{rendered}</p></div>')
        elif style == "LawText":
            output.append(f'<p class="law-text">{rendered}</p>')
        elif text == "GESETZESTEXT":
            output.append('<p class="hero-kicker">Gesetzestext</p>')
        else:
            output.append(f"<p>{rendered}</p>")
    flush_list()
    return "\n".join(output) + "\n"


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    document = Document(SOURCE)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if AI_MARKER.search(full_text):
        raise RuntimeError("KI-/Redaktionsrest im DOCX gefunden.")
    paragraph_numbers = {int(value) for value in re.findall(r"§\s+(\d+)", full_text)}
    if paragraph_numbers != set(range(1, 126)):
        raise RuntimeError(f"Erwartet §§ 1-125, gefunden: {len(paragraph_numbers)} eindeutige Paragrafen.")
    ONLINE.parent.mkdir(parents=True, exist_ok=True)
    ONLINE.write_text(build_online(document), encoding="utf-8")
    print(f"WStG 3.0 online: {ONLINE.relative_to(ROOT)} · §§ 1-125 vollständig")


if __name__ == "__main__":
    main()
