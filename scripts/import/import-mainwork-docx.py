#!/usr/bin/env python3
from __future__ import annotations

import html
import json
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from hashlib import sha256
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_DOCX = Path("/Users/hagen/Desktop/WÖk-Konzepte etc/Kerndokumente/Neuauflage_Buch/Natalie-Weber_Die neue Ordnung des Wohlstands.docx")
OUTPUT = ROOT / "referenz/index.html"
ASSET_DIR = ROOT / "public/assets/imported/woek-main-2026"
PUBLIC_ASSET_PREFIX = "../public/assets/imported/woek-main-2026"
PDF_URL = "../assets/pdf/die-neue-ordnung-des-wohlstands.pdf"


@dataclass
class ImportStats:
    paragraphs: int = 0
    headings: int = 0
    tables: int = 0
    figures: int = 0
    empty_blocks: int = 0


def slugify(value: str) -> str:
    value = value.lower()
    value = value.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "abschnitt"


def file_hash(path: Path) -> str:
    h = sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def iter_blocks(parent: DocumentType) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_image_rel_ids(paragraph: Paragraph) -> list[str]:
    rel_ids = []
    for blip in paragraph._element.xpath(".//*[local-name()='blip']"):
        rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rel_id:
            rel_ids.append(rel_id)
    return rel_ids


def extract_media(docx_path: Path) -> dict[str, str]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    media = {}
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            raw_name = Path(name).name
            target_name = re.sub(r"[^A-Za-z0-9._-]+", "-", raw_name)
            target = ASSET_DIR / target_name
            with archive.open(name) as src, target.open("wb") as dst:
                shutil.copyfileobj(src, dst)
            web_target_name = target_name
            if target.suffix.lower() in {".tif", ".tiff"}:
                png_target = target.with_suffix(".png")
                if png_target.exists():
                    web_target_name = png_target.name
            media[raw_name] = f"{PUBLIC_ASSET_PREFIX}/{web_target_name}"
    return media


def relationship_targets(doc: Document) -> dict[str, str]:
    targets = {}
    for rel_id, rel in doc.part.rels.items():
        target = str(rel.target_ref)
        if target.startswith("media/"):
            targets[rel_id] = Path(target).name
    return targets


def render_table(table: Table, table_no: int) -> str:
    rows = []
    for row in table.rows:
        cells = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
        if cells:
            rows.append(f"<tr>{cells}</tr>")
    return (
        f'<div class="reference-table" id="woek-main-table-{table_no:03d}">'
        f"<table>{''.join(rows)}</table>"
        f"</div>"
    )


def render_page(docx_path: Path) -> dict[str, object]:
    doc = Document(str(docx_path))
    media = extract_media(docx_path)
    rel_targets = relationship_targets(doc)
    stats = ImportStats()
    section_no = 0
    paragraph_no = 0
    table_no = 0
    figure_no = 0
    toc_items = []
    body = []

    source_hash = file_hash(docx_path)

    for block in iter_blocks(doc):
      if isinstance(block, Paragraph):
        text = block.text.strip()
        style = block.style.name if block.style else ""
        image_rel_ids = paragraph_image_rel_ids(block)

        if not text and not image_rel_ids:
            stats.empty_blocks += 1
            continue

        escaped = html.escape(text)
        if style.startswith("toc") and text:
            toc_items.append(text)
            continue

        if style.startswith("Heading"):
            level_match = re.search(r"(\d+)", style)
            level = min(4, max(2, int(level_match.group(1)) + 1 if level_match else 2))
            section_no += 1
            stats.headings += 1
            section_id = f"woek-main-s{section_no:04d}-{slugify(text)[:60]}"
            body.append(f'<h{level} id="{section_id}">{escaped}</h{level}>')
        elif text:
            paragraph_no += 1
            stats.paragraphs += 1
            body.append(f'<p id="woek-main-p{paragraph_no:06d}">{escaped}</p>')

        for rel_id in image_rel_ids:
            raw_name = rel_targets.get(rel_id)
            if raw_name and raw_name in media:
                figure_no += 1
                stats.figures += 1
                src = media[raw_name]
                body.append(
                    f'<figure id="woek-main-fig-{figure_no:04d}" class="reference-figure">'
                    f'<img src="{html.escape(src)}" alt="Die neue Ordnung des Wohlstands, Abbildung {figure_no}">'
                    f'<figcaption>Abbildung {figure_no} aus dem Originaldokument.</figcaption>'
                    f"</figure>"
                )

      elif isinstance(block, Table):
        table_no += 1
        stats.tables += 1
        body.append(render_table(block, table_no))

    toc_html = "\n".join(
        f"<li>{html.escape(item)}</li>"
        for item in toc_items[:400]
    )
    if len(toc_items) > 400:
        toc_html += f"<li>... {len(toc_items) - 400} weitere Inhaltsverzeichniszeilen im Original</li>"

    html_text = f"""<!DOCTYPE html>
<html lang="de">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Die neue Ordnung des Wohlstands - Wirkungsökonomie Online</title>
    <meta name="description" content="Web-Volltext des Grundlagenwerks Die neue Ordnung des Wohlstands von Natalie Weber.">
    <meta name="search_title" content="Die neue Ordnung des Wohlstands">
    <meta name="search_description" content="Volltext des Grundlagenwerks der Wirkungsökonomie mit stabilen Abschnitts- und Absatz-IDs.">
    <meta name="search_section" content="Hauptwerk">
    <meta name="search_type" content="Hauptwerk">
    <link rel="stylesheet" href="../assets/css/style.css?v=20260605-wirkungsraum-stage8">
  </head>
  <body>
    <header class="site-header">
      <a class="brand" href="../index.html" aria-label="Wirkungsökonomie Startseite">
        <span class="brand-mark"><img src="../assets/img/brand/signet.svg" alt="Wirkungsökonomie Logo"></span>
        <span class="brand-name">Wirkungsökonomie</span>
      </a>
      <nav class="site-nav" aria-label="Hauptnavigation">
        <a href="../index.html">Start</a>
        <a href="../verstehen.html">Verstehen</a>
        <a href="../modell.html">Modell</a>
        <a href="../begriffe/">Begriffe</a>
        <a href="../dokumente/">Dokumente</a>
        <a href="../suche.html">Suche</a>
      </nav>
    </header>
    <main class="reference-work" data-pagefind-body>
      <section class="hero compact-hero">
        <p class="hero-kicker">Wirkungsökonomie Online</p>
        <h1>Die neue Ordnung des Wohlstands</h1>
        <p class="hero-subtitle">Begründung und Grundlagen der Wirkungsökonomie. Web-Volltext der bestätigten DOCX-Fassung.</p>
        <p class="notice">Diese Webfassung ist ein technischer Source-Original-Import. Begriffspräzisierungen werden gesondert protokolliert; der Originaltext bleibt zitierfähig.</p>
        <p><a class="button" href="{PDF_URL}">Original-PDF öffnen</a></p>
      </section>
      <aside class="meta-box" id="woek-main-meta">
        <h2>Metadaten</h2>
        <dl>
          <dt>Autorin</dt><dd>Natalie Weber</dd>
          <dt>Dokumenttyp</dt><dd>hauptwerk</dd>
          <dt>Status</dt><dd>source-original</dd>
          <dt>Source-Version</dt><dd>2026.0</dd>
          <dt>Web-Version</dt><dd>2026.1-import</dd>
          <dt>Quelle</dt><dd>Natalie-Weber_Die neue Ordnung des Wohlstands.docx</dd>
          <dt>Source-Hash</dt><dd>{source_hash}</dd>
          <dt>Absätze</dt><dd>{stats.paragraphs}</dd>
          <dt>Überschriften</dt><dd>{stats.headings}</dd>
          <dt>Tabellen</dt><dd>{stats.tables}</dd>
          <dt>Abbildungen</dt><dd>{stats.figures}</dd>
        </dl>
      </aside>
      <section class="content-band" id="woek-main-original-toc">
        <h2>Inhaltsverzeichnis aus dem Originaldokument</h2>
        <ol>{toc_html}</ol>
      </section>
      <article class="article-shell" id="woek-main-fulltext">
        {"".join(body)}
      </article>
      <section class="callout" id="woek-main-discussion-placeholder">
        <h2>Diskurs</h2>

      </section>
    </main>
  </body>
</html>
"""

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(html_text, encoding="utf-8")
    return {
        "documentId": "woek-main-2026",
        "sourceFile": "Natalie-Weber_Die neue Ordnung des Wohlstands.docx",
        "sourceVersion": "2026.0",
        "webVersion": "2026.1-import",
        "contentState": "source-original",
        "sourceHash": source_hash,
        "route": "/referenz/",
        "stats": stats.__dict__,
        "assetDir": str(ASSET_DIR.relative_to(ROOT)),
    }


def main() -> None:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not src.exists():
        raise SystemExit(f"Missing mainwork DOCX: {src}")
    result = render_page(src)
    data_dir = ROOT / "public/data"
    data_dir.mkdir(parents=True, exist_ok=True)
    (data_dir / "mainwork-import.json").write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Wrote {OUTPUT.relative_to(ROOT)} with {result['stats']['paragraphs']} paragraphs and {result['stats']['figures']} figures.")


if __name__ == "__main__":
    main()
