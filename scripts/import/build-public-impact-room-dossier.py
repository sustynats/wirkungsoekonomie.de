#!/usr/bin/env python3
"""Build the public impact room long dossier from the editorial DOCX source."""

from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(
    "/Users/hagen/Downloads/Dossier_Oeffentlicher_Wirkungsraum_Wellen_Tiefe_WOeK_100Seiten_Arbeitsfassung.docx"
)
SOURCE_COPY = ROOT / "docs/oeffentlicher-wirkungsraum/source/woek_dossier_oeffentlicher_wirkungsraum_wellen_tiefe_v0_1.docx"
PDF_TARGET = ROOT / "assets/downloads/woek_dossier_oeffentlicher_wirkungsraum_wellen_tiefe_v0_1.pdf"
PAGE_TARGET = ROOT / "oeffentlicher-wirkungsraum/dossier-stein-wellen-tiefe/index.html"

TITLE = "Der öffentliche Wirkungsraum: Stein, Wellen und Tiefe"
SUBTITLE = (
    "Ein Langdossier über Debatten, Resonanz, Aufmerksamkeit und demokratische Resilienz."
)
CANONICAL = "https://wirkungsoekonomie.de/oeffentlicher-wirkungsraum/dossier-stein-wellen-tiefe/"
DATE = "2026-06-07"
BASE = "../../"


def block_items(parent):
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def set_cell_text(cell, value):
    cell.text = value


def set_style_font(style, name, size=None, bold=None, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor(*color)


def apply_woek_docx_style(src: Path, target: Path):
    target.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(src)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    set_style_font(styles["Normal"], "Arial", 10.5, False, (30, 30, 30))
    for name, size in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)]:
        if name in styles:
            set_style_font(styles[name], "Georgia", size, True, (10, 16, 32))
    if "Quote" in styles:
        set_style_font(styles["Quote"], "Arial", 11, False, (36, 96, 72))

    header = section.header
    if header.paragraphs:
        header.paragraphs[0].text = "Wirkungsökonomie · Öffentlicher Wirkungsraum"
        header.paragraphs[0].style = styles["Normal"]
    footer = section.footer
    if footer.paragraphs:
        footer.paragraphs[0].text = "Dossier · Stein, Wellen und Tiefe · wirkungsoekonomie.de"
        footer.paragraphs[0].style = styles["Normal"]

    # Normalize empty leading paragraphs but keep the editorial structure intact.
    for para in doc.paragraphs:
        if para.text.strip():
            fmt = para.paragraph_format
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.15
    doc.save(target)


def convert_pdf(source: Path, target: Path):
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    result = subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target.parent),
            str(source),
        ],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    generated = source.with_suffix(".pdf")
    if generated != target and generated.exists():
        generated.replace(target)
    if result.returncode != 0 or not target.exists():
        raise RuntimeError(
            f"PDF conversion failed: {result.returncode}\nSTDOUT={result.stdout}\nSTDERR={result.stderr}"
        )


def slugify(value: str) -> str:
    value = value.lower()
    replacements = {
        "ä": "ae",
        "ö": "oe",
        "ü": "ue",
        "ß": "ss",
        "&": "und",
    }
    for src, dst in replacements.items():
        value = value.replace(src, dst)
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "abschnitt"


LINKS = {
    "Öffentlicher Wirkungsraum": "../",
    "Debatten-Kompass": "../../wirkungsradar/",
    "Debattenkarten": "../../wirkungsradar/debattenkarten/",
    "Resonanz-Kompass": "../../wirkungsradar/resonanz-kompass/",
    "Agenda-Radar": "../../wirkungsradar/agenda-radar/",
    "Ursachen-Navigator": "../../wirkungsradar/ursachen-navigator/",
    "Resilienz-Prinzipien": "../../wirkungsradar/resilienz-prinzipien/",
    "Wirkungspotenzial": "../../begriffe/wirkungspotenzial/",
    "Wirkungsrisiko": "../../begriffe/wirkungsrisiko/",
    "Wirkpfad": "../../begriffe/wirkpfad/",
    "Resonanzraum": "../../begriffe/resonanzraum/",
    "Narrativ": "../../begriffe/narrativ/",
    "Wirkungsrückkopplung": "../../begriffe/wirkungsrueckkopplung/",
    "Wirkungsblindheit": "../../begriffe/wirkungsblindheit/",
    "Wirkungswahrheit": "../../begriffe/wirkungswahrheit/",
    "Wirkungsabwehr": "../../begriffe/wirkungsabwehr/",
    "Dissonanzrationalisierung": "../../begriffe/dissonanzrationalisierung/",
    "Kognitive Dissonanz": "../../begriffe/kognitive-dissonanz/",
    "Mensch, Planet und Demokratie": "../../begriffe/mensch-planet-demokratie/",
    "positive Netto-Wirkung": "../../begriffe/positive-netto-wirkung/",
    "Wirkung": "../../begriffe/wirkung/",
}


def link_text(text: str, counts: defaultdict[str, int]) -> str:
    escaped = html.escape(text)
    # Link specific multi-word terms first. Cap repeated links to keep the long dossier readable.
    for label, href in sorted(LINKS.items(), key=lambda item: len(item[0]), reverse=True):
        if counts[label] >= 6:
            continue
        pattern = re.compile(rf"(?<![\w>]){re.escape(html.escape(label))}(?![\w<])")

        def repl(match):
            if counts[label] >= 6:
                return match.group(0)
            counts[label] += 1
            return f'<a class="text-link" href="{href}">{match.group(0)}</a>'

        escaped = pattern.sub(repl, escaped)
    return escaped


def cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def render_table(table, table_index: int, link_counts) -> str:
    rows = []
    for r_index, row in enumerate(table.rows):
        cells = []
        tag = "th" if r_index == 0 else "td"
        for cell in row.cells:
            text = cell_text(cell)
            cells.append(f"<{tag}>{link_text(text, link_counts)}</{tag}>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return (
        f'<div class="dossier-table-wrap" id="tabelle-{table_index:03d}">'
        f'<table class="dossier-table">{"".join(rows)}</table></div>'
    )


def paragraph_level(style_name: str) -> int | None:
    match = re.match(r"Heading\s+([1-6])", style_name or "")
    return int(match.group(1)) if match else None


def render_online_page(source: Path, target: Path):
    doc = Document(source)
    link_counts = defaultdict(int)
    body_parts = []
    toc = []
    slug_counts = defaultdict(int)
    table_index = 0
    skip_empty = 0

    for block in block_items(doc):
        if hasattr(block, "rows"):
            table_index += 1
            body_parts.append(render_table(block, table_index, link_counts))
            continue

        text = block.text.strip()
        if not text:
            skip_empty += 1
            continue
        style = block.style.name if block.style else ""

        # The first normal paragraph is a cover/title block. The HTML hero carries it cleaner.
        if text.startswith("Der öffentliche Wirkungsraum") and "Arbeitsfassung" in text:
            continue

        level = paragraph_level(style)
        if level:
            tag = "h2" if level == 1 else "h3" if level == 2 else "h4"
            slug = slugify(text)
            slug_counts[slug] += 1
            if slug_counts[slug] > 1:
                slug = f"{slug}-{slug_counts[slug]}"
            toc.append((level, text, slug))
            body_parts.append(f'<{tag} id="{slug}">{link_text(text, link_counts)}</{tag}>')
            continue

        if style == "Quote":
            body_parts.append(f'<blockquote>{link_text(text, link_counts)}</blockquote>')
        elif "List" in style:
            body_parts.append(f'<p class="dossier-list-item">{link_text(text, link_counts)}</p>')
        else:
            body_parts.append(f'<p>{link_text(text, link_counts)}</p>')

    toc_items = []
    for level, text, slug in toc:
        if level == 1:
            toc_items.append(f'<li><a href="#{slug}">{html.escape(text)}</a></li>')
    toc_html = "\n".join(toc_items)

    body_html = "\n".join(body_parts)
    page = f"""<!doctype html>
<html lang="de">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(TITLE)} | Wirkungsökonomie</title>
  <meta name="description" content="{html.escape(SUBTITLE)}">
  <meta name="search_title" content="{html.escape(TITLE)}">
  <meta name="search_type" content="Dossier">
  <meta name="search_tags" content="Öffentlicher Wirkungsraum, Debatten-Kompass, Resonanz-Kompass, Agenda-Radar, Ursachen-Navigator, Resilienz">
  <link rel="canonical" href="{CANONICAL}">
  <link rel="stylesheet" href="{BASE}assets/css/style.css?v=20260612-journal-mobile-fix">
  <style>
    .dossier-layout {{ max-width: 1080px; margin: 0 auto; }}
    .dossier-prose {{ max-width: 920px; margin: 0 auto; }}
    .dossier-prose p {{ font-size: 1.03rem; line-height: 1.72; }}
    .dossier-prose blockquote {{ border-left: 4px solid #2f7f5f; margin: 1.6rem 0; padding: 1rem 1.25rem; background: rgba(47,127,95,.06); font-weight: 700; }}
    .dossier-list-item {{ padding-left: 1.4rem; position: relative; }}
    .dossier-list-item::before {{ content: "•"; position: absolute; left: .25rem; color: #2f7f5f; font-weight: 800; }}
    .dossier-table-wrap {{ overflow-x: auto; margin: 1.5rem 0; border: 1px solid rgba(20,20,30,.14); border-radius: 8px; background: #fffdfa; }}
    .dossier-table {{ width: 100%; border-collapse: collapse; min-width: 680px; }}
    .dossier-table th, .dossier-table td {{ padding: .75rem .85rem; border-bottom: 1px solid rgba(20,20,30,.12); vertical-align: top; text-align: left; line-height: 1.45; }}
    .dossier-table th {{ color: #144f3a; font-weight: 800; background: rgba(47,127,95,.08); }}
    .dossier-meta-row {{ display:flex; flex-wrap:wrap; gap:.75rem; margin-top: 1rem; color:#5d5a55; font-weight:700; }}
  </style>
</head>
<body>
  <header class="site-header" data-search-exclude>
    <a class="brand" href="{BASE}index.html"><span class="brand-mark"><img src="{BASE}assets/img/brand/signet.svg" alt="Wirkungsökonomie Logo"></span><span class="brand-name">Wirkungsökonomie</span></a>
    <button class="nav-toggle" type="button" aria-label="Menü öffnen" aria-expanded="false" aria-controls="site-nav"><span class="nav-toggle-icon" aria-hidden="true">☰</span><span class="sr-only">Menü</span></button>
    <nav class="site-nav" id="site-nav" aria-label="Hauptnavigation" data-search-exclude></nav>
  </header>
  <main>
    <section class="hero portal-hero">
      <div class="hero-content">
        <nav class="breadcrumb" aria-label="Breadcrumb"><a href="{BASE}index.html">Start</a> / <a href="{BASE}oeffentlicher-wirkungsraum/">Öffentlicher Wirkungsraum</a></nav>
        <p class="hero-kicker">Langdossier</p>
        <h1>{html.escape(TITLE)}</h1>
        <p class="hero-subtitle">{html.escape(SUBTITLE)}</p>
        <p>Die meisten Menschen diskutieren über den Stein. Dieses Dossier analysiert die Wellen - und die Tiefe, aus der öffentliche Wirkung möglich wird.</p>
        <p class="dossier-meta-row"><span>Arbeitsfassung</span><span>Stand: {DATE}</span><span>Online lesbar mit Abschnittsankern</span></p>
        <div class="hero-actions no-print">
          <a class="btn btn-primary" href="#onlinefassung">Online lesen</a>
          <a class="btn btn-secondary" href="{BASE}assets/downloads/{PDF_TARGET.name}">PDF herunterladen</a>
          <a class="btn btn-secondary" href="{BASE}oeffentlicher-wirkungsraum/">Wirkungsraum öffnen</a>
        </div>
      </div>
    </section>
    <section class="section dossier-layout" aria-labelledby="toc-title">
      <details class="card" open>
        <summary><strong id="toc-title">Inhaltsverzeichnis</strong></summary>
        <ol class="toc-list">{toc_html}</ol>
      </details>
    </section>
    <section class="section" id="onlinefassung" aria-labelledby="onlinefassung-title">
      <div class="dossier-prose">
        <p class="hero-kicker">Onlinefassung</p>
        <h2 id="onlinefassung-title">Dossier lesen</h2>
        {body_html}
      </div>
    </section>
  </main>
  <script src="{BASE}assets/js/main.js?v=20260612-journal-mobile-fix"></script>
</body>
</html>
"""
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(page, encoding="utf-8")
    print(f"Wrote online dossier: {target.relative_to(ROOT)}")
    print(f"Headings: {len(toc)}, tables: {table_index}, skipped empty paragraphs: {skip_empty}")


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    apply_woek_docx_style(SOURCE, SOURCE_COPY)
    convert_pdf(SOURCE_COPY, PDF_TARGET)
    render_online_page(SOURCE, PAGE_TARGET)
    print(f"Wrote source DOCX: {SOURCE_COPY.relative_to(ROOT)}")
    print(f"Wrote PDF: {PDF_TARGET.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
