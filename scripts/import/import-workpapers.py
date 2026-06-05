#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import html
import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path.cwd()
ORIGINALS_DIR = ROOT / "public" / "downloads" / "originals"
DOCUMENTS_DIR = ROOT / "dokumente"
DATA_FILE = ROOT / "public" / "data" / "workpaper-imports.json"

WORKPAPERS = [
    {
        "title": "WStG Oktober 2025",
        "slug": "wstg-oktober-2025",
        "documentType": "gesetzesentwurf",
        "status": "gesetzesentwurf",
        "source": "assets/pdf/wirkungssteuergesetz-wstg-oktober-2025.pdf",
        "originalName": "WStG_Oktober2025.pdf",
    },
    {
        "title": "Grundlagenpapier Wirkungsökonomie WÖk",
        "slug": "grundlagenpapier-wirkungsoekonomie-woek",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Grundlagenpapier-Wirkungsökonomie WÖk.pdf",
        "originalName": "Grundlagenpapier-Wirkungsökonomie WÖk.pdf",
    },
    {
        "title": "WÖk Master Items final v1.2",
        "slug": "woek-master-items-final-v1-2",
        "documentType": "register",
        "status": "register",
        "source": "/Users/hagen/Desktop/Alter-Schreibtsch/Schreibtisch3/WOeK_Master_Items_final_v1.2.pdf",
        "originalName": "WOeK_Master_Items_final_v1.2.pdf",
    },
    {
        "title": "Technische Leitlinien WUStG Vollversion Extended v2",
        "slug": "technische-leitlinien-wustg-v2",
        "documentType": "technische-leitlinie",
        "status": "technische-leitlinie",
        "source": "/Users/hagen/Downloads/Technische_Leitlinien_WUStG_Vollversion_Extended_v2.pdf",
        "originalName": "Technische_Leitlinien_WUStG_Vollversion_Extended_v2.pdf",
    },
    {
        "title": "Beispiel Apfel Wirkungssteuer Bonusregel",
        "slug": "beispiel-apfel-wirkungssteuer-bonusregel",
        "documentType": "beispiel",
        "status": "beispielrechnung",
        "source": "/Users/hagen/Downloads/Beispiel_Apfel_Wirkungssteuer_Bonusregel.pdf",
        "originalName": "Beispiel_Apfel_Wirkungssteuer_Bonusregel.pdf",
    },
    {
        "title": "Wirkungsrat Konzept",
        "slug": "wirkungsrat-konzept",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Wirkungsrat_Konzept.pdf",
        "originalName": "Wirkungsrat_Konzept.pdf",
    },
    {
        "title": "Whitepaper T-SROI",
        "slug": "whitepaper-t-sroi",
        "documentType": "whitepaper",
        "status": "arbeitspapier",
        "source": "/Users/hagen/Downloads/Whitepaper-T-SROI.pdf",
        "originalName": "Whitepaper-T-SROI.pdf",
    },
    {
        "title": "Wirkungsökonomie in der Lieferkette",
        "slug": "wirkungsoekonomie-in-der-lieferkette",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Wirkungsökonomie in der Lieferkette.pdf",
        "originalName": "Wirkungsökonomie in der Lieferkette.pdf",
    },
    {
        "title": "Systemmodell der Wirkungsökonomie",
        "slug": "systemmodell-der-wirkungsoekonomie",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Systemmodell-der-Wirkungsökonomie.pdf",
        "originalName": "Systemmodell-der-Wirkungsökonomie.pdf",
    },
    {
        "title": "WP Produkte",
        "slug": "wp-produkte",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "assets/pdf/working-paper-produktbesteuerung-durch-wirkung.pdf",
        "originalName": "WP_Produkte.pdf",
    },
    {
        "title": "WP Einkommen",
        "slug": "wp-einkommen",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "assets/pdf/whitepaper-wirkungseinkommen.pdf",
        "originalName": "WP_Einkommen.pdf",
    },
    {
        "title": "WP Wohnungsmarkt",
        "slug": "wp-wohnungsmarkt",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "assets/pdf/working-paper-wohnungsmarkt.pdf",
        "originalName": "WP_Wohnungsmarkt_.pdf",
    },
    {
        "title": "Wenn Maschinen arbeiten",
        "slug": "wenn-maschinen-arbeiten",
        "documentType": "arbeitspapier",
        "status": "arbeitspapier",
        "source": "assets/pdf/wenn-maschinen-arbeiten.pdf",
        "originalName": "Wenn Maschinen arbeiten.pdf",
    },
    {
        "title": "Leitbild für Mensch, Planet und Demokratie",
        "slug": "leitbild-mensch-planet-demokratie",
        "documentType": "leitbild",
        "status": "arbeitspapier",
        "source": "assets/pdf/leitbild-mensch-planet-demokratie.pdf",
        "originalName": "Leitbild für Mensch Planet und Demokratie.pdf",
    },
    {
        "title": "Minifest Wirkungsökonomie",
        "slug": "minifest-wirkungsoekonomie",
        "documentType": "manifest",
        "status": "archiv",
        "source": "/Users/hagen/Downloads/anschreiben1_kampagne/Minifest_Wirkungsoekonomie.pdf",
        "originalName": "Minifest_Wirkungsoekonomie.pdf",
    },
    {
        "title": "WÖk-Manifest",
        "slug": "woek-manifest",
        "documentType": "manifest",
        "status": "archiv",
        "source": "/Users/hagen/Downloads/WÖk-Manifest.pdf",
        "originalName": "WÖk-Manifest.pdf",
    },
    {
        "title": "WÖK-Partei",
        "slug": "woek-partei",
        "documentType": "parteiprogramm",
        "status": "archiv",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/WÖK-Partei.pdf",
        "originalName": "WÖK-Partei.pdf",
    },
    {
        "title": "NATS WÖk allgemein",
        "slug": "nats-woek-allgemein",
        "documentType": "praesentation",
        "status": "archiv",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/NATS_WÖk@allgemein.pdf",
        "originalName": "NATS_WÖk@allgemein.pdf",
    },
    {
        "title": "Beispiel Konzern",
        "slug": "beispiel-konzern",
        "documentType": "beispiel",
        "status": "beispielrechnung",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Beispiel-Konzern.pdf",
        "originalName": "Beispiel-Konzern.pdf",
    },
    {
        "title": "FAZ-Beitrag",
        "slug": "faz-beitrag",
        "documentType": "artikel",
        "status": "archiv",
        "source": "/Users/hagen/Desktop/WÖk-Konzepte etc/Veröffentlichungen/FAZ-Beitrag.docx",
        "originalName": "FAZ-Beitrag.docx",
    },
]


def digest(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def escape(value: str) -> str:
    return html.escape(value or "", quote=True)


def paragraph_id(slug: str, index: int) -> str:
    return f"{slug}-p{index:04d}"


def section_id(slug: str, index: int) -> str:
    return f"{slug}-s{index:04d}"


def extract_pdf(path: Path) -> tuple[list[str], list[str]]:
    reader = PdfReader(str(path))
    paragraphs: list[str] = []
    issues: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            issues.append(f"Seite {page_index}: Text konnte nicht extrahiert werden ({exc}).")
            continue
        cleaned = re.sub(r"\n{3,}", "\n\n", text.strip())
        if not cleaned:
            issues.append(f"Seite {page_index}: kein extrahierbarer Text.")
            continue
        paragraphs.append(f"Seite {page_index}")
        paragraphs.extend([part.strip() for part in re.split(r"\n\s*\n|(?<=\.)\n(?=[A-ZÄÖÜ])", cleaned) if part.strip()])
    return paragraphs, issues


def extract_docx(path: Path) -> tuple[list[str], list[str]]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    issues: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append(f"Tabelle {table_index}")
            paragraphs.extend(rows)
    return paragraphs, issues


def render_page(item: dict, paragraphs: list[str], issues: list[str], original_href: str, source_hash: str) -> str:
    body_parts = []
    current_section_id = section_id(item["slug"], 1)
    for idx, paragraph in enumerate(paragraphs, start=1):
        if re.fullmatch(r"(Seite|Tabelle)\s+\d+", paragraph):
            current_section_id = section_id(item["slug"], idx)
            body_parts.append(
                f'<h2 id="{current_section_id}" data-document-id="{escape(item["slug"])}" '
                f'data-section-id="{current_section_id}" data-version="2026.1-import" '
                f'data-content-hash="{hashlib.sha256(paragraph.encode("utf-8")).hexdigest()[:16]}">{escape(paragraph)}</h2>'
            )
        else:
            pid = paragraph_id(item["slug"], idx)
            body_parts.append(
                f'<p id="{pid}" data-document-id="{escape(item["slug"])}" data-section-id="{current_section_id}" '
                f'data-paragraph-id="{pid}" data-version="2026.1-import" '
                f'data-content-hash="{hashlib.sha256(paragraph.encode("utf-8")).hexdigest()[:16]}">{escape(paragraph)}</p>'
            )
    issue_html = ""
    if issues:
        issue_html = "<section class=\"callout\"><h2>Konvertierungshinweise</h2><ul>" + "".join(
            f"<li>{escape(issue)}</li>" for issue in issues
        ) + "</ul></section>"
    return f"""<!DOCTYPE html>
<html lang="de">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{escape(item["title"])} - Wirkungsökonomie Online</title>
    <meta name="description" content="Webfassung und Originaldatei: {escape(item["title"])}.">
    <meta name="search_title" content="{escape(item["title"])}">
    <meta name="search_description" content="Webfassung eines Arbeitspapiers der Wirkungsökonomie mit Originaldatei.">
    <meta name="search_section" content="Dokumente">
    <meta name="search_type" content="{escape(item["documentType"])}">
    <link rel="stylesheet" href="../../assets/css/style.css?v=20260605-wirkungsraum-stage8">
  </head>
  <body>
    <header class="site-header">
      <a class="brand" href="../../index.html" aria-label="Wirkungsökonomie Startseite">
        <span class="brand-mark"><img src="../../assets/img/brand/signet.svg" alt="Wirkungsökonomie Logo"></span>
        <span class="brand-name">Wirkungsökonomie</span>
      </a>
      <nav class="site-nav" id="site-nav" aria-label="Hauptnavigation">
        <a href="../../index.html">Start</a>
        <a href="../../referenz/">Referenz</a>
        <a href="../../begriffe/">Begriffe</a>
        <a href="../../dokumente/">Dokumente</a>
        <a href="../../suche.html">Suche</a>
      </nav>
    </header>
    <main class="reference-work" data-pagefind-body>
      <article class="article-shell">
        <nav class="breadcrumb"><a href="../">Dokumente</a> / {escape(item["title"])}</nav>
        <h1>{escape(item["title"])}</h1>
        <p class="lead">Webfassung aus der gelieferten Originaldatei. Der Originaltext bleibt über die Originaldatei zitierbar.</p>
        <p><a class="button" href="{escape(original_href)}">Originaldatei öffnen</a></p>
        <section class="meta-box">
          <h2>Metadaten</h2>
          <dl>
            <dt>Dokumenttyp</dt><dd>{escape(item["documentType"])}</dd>
            <dt>Status</dt><dd>{escape(item["status"])}</dd>
            <dt>Source-Version</dt><dd>2026.0</dd>
            <dt>Web-Version</dt><dd>2026.1-import</dd>
            <dt>Reviewstatus</dt><dd>partially-reviewed</dd>
            <dt>Terminologiebasis</dt><dd>WOeK_Begriffsleitfaden_fuehrend_v1.0.md</dd>
            <dt>Originaldatei</dt><dd>{escape(item["originalName"])}</dd>
            <dt>Source-Hash</dt><dd>{escape(source_hash)}</dd>
            <dt>Absätze/Textblöcke</dt><dd>{len(paragraphs)}</dd>
          </dl>
        </section>
        <section class="callout">
          <h2>Importstatus</h2>
          <p>Diese Webfassung ist ein technischer Volltextimport. Layout, Fußnotenpositionen, komplexe Tabellen und eingebettete Grafiken können vom Original abweichen; die Originaldatei bleibt die zitierfähige Fassung.</p>
          <p>Diskurs zu einzelnen Abschnitten wird in Phase 2 aktiviert; die Abschnitts- und Absatz-IDs sind vorbereitet.</p>
        </section>
        {issue_html}
        <section>
          <h2>Web-Volltext</h2>
          {''.join(body_parts)}
        </section>
      </article>
    </main>
  </body>
</html>
"""


def render_library(imports: list[dict]) -> str:
    cards = [
        {
            "title": "Die neue Ordnung des Wohlstands",
            "slug": "../referenz/",
            "type": "hauptwerk",
            "status": "source-original",
            "description": "Web-Volltext der bestätigten DOCX-Fassung mit Original-PDF.",
            "original": "../assets/pdf/die-neue-ordnung-des-wohlstands.pdf",
        }
    ]
    for item in imports:
        cards.append(
            {
                "title": item["title"],
                "slug": f'{item["slug"]}/',
                "type": item["documentType"],
                "status": item["status"],
                "description": "Webfassung mit Originaldatei.",
                "original": item["originalUrl"],
            }
        )
    card_html = "\n".join(
        f"""<article class="info-card">
          <p class="meta-line">{escape(card["type"])} · {escape(card["status"])}</p>
          <h3><a href="{escape(card["slug"])}">{escape(card["title"])}</a></h3>
          <p>{escape(card["description"])}</p>
          <p><a href="{escape(card["original"])}">Originaldatei öffnen</a></p>
        </article>"""
        for card in cards
    )
    return f"""<!DOCTYPE html>
<html lang="de">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Dokumente der Wirkungsökonomie Online</title>
    <meta name="description" content="Dokumentenbibliothek mit Hauptwerk, Arbeitspapieren und Originaldateien der Wirkungsökonomie.">
    <meta name="search_title" content="Dokumente der Wirkungsökonomie Online">
    <meta name="search_description" content="Dokumentenbibliothek mit Hauptwerk, Arbeitspapieren und Originaldateien der Wirkungsökonomie.">
    <meta name="search_section" content="Dokumente">
    <meta name="search_type" content="Dokumentenbibliothek">
    <link rel="stylesheet" href="../assets/css/style.css?v=20260605-wirkungsraum-stage8">
  </head>
  <body>
    <header class="site-header">
      <a class="brand" href="../index.html" aria-label="Wirkungsökonomie Startseite">
        <span class="brand-mark"><img src="../assets/img/brand/signet.svg" alt="Wirkungsökonomie Logo"></span>
        <span class="brand-name">Wirkungsökonomie</span>
      </a>
      <nav class="site-nav" id="site-nav" aria-label="Hauptnavigation">
        <a href="../index.html">Start</a>
        <a href="../referenz/">Referenz</a>
        <a href="../begriffe/">Begriffe</a>
        <a href="../dokumente/">Dokumente</a>
        <a href="../suche.html">Suche</a>
      </nav>
    </header>
    <main class="section" data-pagefind-body>
      <section class="hero compact-hero">
        <p class="hero-kicker">Wirkungsökonomie Online</p>
        <h1>Dokumente der Wirkungsökonomie Online</h1>
        <p class="hero-subtitle">Hauptwerk, Arbeitspapiere, Leitlinien, Beispiele, Manifeste und Archivdokumente als Webfassung mit Originaldatei.</p>
      </section>
      <section class="card-grid">
        {card_html}
      </section>
    </main>
  </body>
</html>
"""


def main() -> None:
    ORIGINALS_DIR.mkdir(parents=True, exist_ok=True)
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    imports = []
    for item in WORKPAPERS:
        source = Path(item["source"])
        if not source.exists():
            raise FileNotFoundError(source)
        target = ORIGINALS_DIR / item["originalName"]
        shutil.copy2(source, target)
        source_hash = digest(target)
        if target.suffix.lower() == ".pdf":
            paragraphs, issues = extract_pdf(target)
        elif target.suffix.lower() == ".docx":
            paragraphs, issues = extract_docx(target)
        else:
            paragraphs, issues = [], [f"Dateityp {target.suffix} wird noch nicht extrahiert."]
        if not paragraphs:
            issues.append("Kein Volltext extrahiert; bitte bessere Quelldatei bereitstellen.")
        doc_dir = DOCUMENTS_DIR / item["slug"]
        doc_dir.mkdir(parents=True, exist_ok=True)
        original_href = f"../../public/downloads/originals/{item['originalName']}"
        page = render_page(item, paragraphs, issues, original_href, source_hash)
        (doc_dir / "index.html").write_text(page, encoding="utf-8")
        imports.append(
            {
                **item,
                "sourceHash": source_hash,
                "originalUrl": f"../public/downloads/originals/{item['originalName']}",
                "route": f"/dokumente/{item['slug']}/",
                "blocks": len(paragraphs),
                "issues": issues,
            }
        )
    (DOCUMENTS_DIR / "index.html").write_text(render_library(imports), encoding="utf-8")
    DATA_FILE.parent.mkdir(parents=True, exist_ok=True)
    DATA_FILE.write_text(
        json.dumps({"generatedAt": datetime.now(timezone.utc).isoformat(), "documents": imports}, ensure_ascii=False, indent=2)
        + "\n",
        encoding="utf-8",
    )
    print(f"Imported {len(imports)} workpapers.")


if __name__ == "__main__":
    main()
