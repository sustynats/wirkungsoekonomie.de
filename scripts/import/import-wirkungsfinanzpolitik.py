#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import html
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET


def _ensure_docx_runtime() -> None:
    try:
        import docx  # noqa: F401
        return
    except ModuleNotFoundError:
        bundled_python = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "bin" / "python3"
        if bundled_python.exists() and Path(sys.executable).resolve() != bundled_python.resolve():
            os.execv(str(bundled_python), [str(bundled_python), *sys.argv])
        raise


_ensure_docx_runtime()

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path("/Users/hagen/Downloads/Wirkungsfinanzpolitik_Aufsatz_WOeK_v2_MMT_PublicPurpose.docx")
TEMPLATE = ROOT / "assets" / "downloads" / "woek_publikationsstandard_detailkonzepte_dossiers_v0_3.docx"
PUBLIC_PDF_NAME = "wirkungsfinanzpolitik-aufsatz-woek-v2-mmt-public-purpose.pdf"
PUBLIC_PDF_TARGET = ROOT / "public" / "downloads" / "originals" / PUBLIC_PDF_NAME
INTERNAL_DOCX_TARGET = ROOT / "content" / "internal-documents" / "wirkungsfinanzpolitik" / "wirkungsfinanzpolitik-aufsatz-woek-ci.docx"
DOCUMENT_DIR = ROOT / "dokumente" / "wirkungsfinanzpolitik"
LIBRARY_DIR = ROOT / "bibliothek" / "wirkungsfinanzpolitik"
SHELL_PAGE = ROOT / "bibliothek" / "arbeitspapier-doppelte-wesentlichkeit-impact-controlling" / "index.html"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W = f"{{{NS['w']}}}"

TITLE = "Von der Schuldenfrage zur Wirkungsfinanzpolitik"
SUBTITLE = "Öffentliche Finanzen, Staatsschulden und positive Netto-Wirkung aus Sicht der Wirkungsökonomie"
DESCRIPTION = (
    "Arbeitsfassung zur Wirkungsfinanzpolitik: öffentliche Einnahmen, Ausgaben, Kredite, "
    "Zinsen, Steuern und Investitionen werden nach ihrer Netto-Wirkung auf Mensch, Planet "
    "und Demokratie bewertet. Die v2-Fassung ordnet MMT, Functional Finance und Public "
    "Purpose als Anschlussstellen ein."
)
STATUS = "Arbeitsfassung / Entwurf"
STAND = "11. Juni 2026"
SOURCE_VERSION = "2026.0-v2-mmt-public-purpose"
WEB_VERSION = "2026.2-webimport"


def digest(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def esc(value: str) -> str:
    return html.escape(value or "", quote=True)


def slugify(value: str, used: set[str]) -> str:
    slug = value.lower()
    slug = slug.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    slug = re.sub(r"[^a-z0-9]+", "-", slug).strip("-")
    slug = slug or "abschnitt"
    base = slug
    counter = 2
    while slug in used:
        slug = f"{base}-{counter}"
        counter += 1
    used.add(slug)
    return slug


def para_text(para: ET.Element) -> str:
    return "".join(t.text or "" for t in para.findall(".//w:t", NS)).strip()


def para_style(para: ET.Element) -> str:
    ppr = para.find("w:pPr", NS)
    if ppr is None:
        return "Normal"
    style = ppr.find("w:pStyle", NS)
    if style is None:
        return "Normal"
    return style.attrib.get(f"{W}val", "Normal")


def table_rows(table: ET.Element) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.findall("w:tr", NS):
        cells: list[str] = []
        for cell in row.findall("w:tc", NS):
            pieces = [para_text(p) for p in cell.findall("w:p", NS)]
            cells.append(" ".join(piece for piece in pieces if piece))
        if any(cells):
            rows.append(cells)
    return rows


def read_docx_blocks(path: Path) -> list[dict]:
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    body = root.find("w:body", NS)
    if body is None:
        return []
    blocks: list[dict] = []
    for child in body:
        if child.tag == f"{W}p":
            text = para_text(child)
            if text:
                blocks.append({"type": "paragraph", "style": para_style(child), "text": text})
        elif child.tag == f"{W}tbl":
            rows = table_rows(child)
            if rows:
                blocks.append({"type": "table", "rows": rows})
    return blocks


def split_source(blocks: list[dict]) -> tuple[list[dict], list[dict]]:
    intro: list[dict] = []
    body: list[dict] = []
    skipping_toc = False
    reached_body = False
    for block in blocks:
        if not reached_body:
            text = block.get("text", "")
            if block["type"] == "paragraph" and block.get("style") == "Heading1" and text in {"Inhaltsverzeichnis", "Inhaltsübersicht"}:
                skipping_toc = True
                continue
            if skipping_toc:
                if block["type"] == "paragraph" and block.get("style") == "Heading1" and text == "Abstract":
                    reached_body = True
                    body.append(block)
                continue
            intro.append(block)
            continue
        body.append(block)
    return intro, body


def render_body(blocks: list[dict]) -> tuple[str, list[tuple[str, str]]]:
    used: set[str] = set()
    parts: list[str] = []
    toc: list[tuple[str, str]] = []
    list_open = False
    section = "intro"
    paragraph_index = 0

    def close_list() -> None:
        nonlocal list_open
        if list_open:
            parts.append("</ul>")
            list_open = False

    for block in blocks:
        if block["type"] == "table":
            close_list()
            rows = block["rows"]
            parts.append('<div class="table-wrap" role="region" tabindex="0"><table class="data-table">')
            if rows:
                parts.append("<thead><tr>" + "".join(f"<th>{esc(cell)}</th>" for cell in rows[0]) + "</tr></thead>")
                parts.append("<tbody>")
                for row in rows[1:]:
                    parts.append("<tr>" + "".join(f"<td>{esc(cell)}</td>" for cell in row) + "</tr>")
                parts.append("</tbody>")
            parts.append("</table></div>")
            continue

        style = block.get("style", "Normal")
        text = block["text"]

        if style == "ListBullet":
            if not list_open:
                parts.append("<ul>")
                list_open = True
            paragraph_index += 1
            pid = f"wp-{paragraph_index:04d}"
            parts.append(
                f'<li id="{pid}" data-section-id="{esc(section)}" data-paragraph-id="{pid}">{esc(text)}</li>'
            )
            continue

        close_list()

        if style in {"Heading1", "Heading2"}:
            hid = slugify(text, used)
            section = hid
            toc.append((hid, text))
            tag = "h2" if style == "Heading1" else "h3"
            parts.append(
                f'<{tag} id="{hid}" data-section-id="{hid}">{esc(text)} '
                f'<a class="cite-anchor no-print" href="#{hid}" aria-label="Zitierlink zu diesem Abschnitt">#</a></{tag}>'
            )
            continue

        paragraph_index += 1
        pid = f"wp-{paragraph_index:04d}"
        if style in {"Callout", "Kernsatz"}:
            parts.append(
                f'<blockquote id="{pid}" data-section-id="{esc(section)}" data-paragraph-id="{pid}">{esc(text)}</blockquote>'
            )
        else:
            parts.append(
                f'<p id="{pid}" data-section-id="{esc(section)}" data-paragraph-id="{pid}">{esc(text)}</p>'
            )

    close_list()
    return "\n".join(parts), toc


def page_shell_prefix_suffix(relative_depth: int) -> tuple[str, str]:
    page = SHELL_PAGE.read_text(encoding="utf-8")
    header_start = page.index('    <header class="site-header"')
    main_start = page.index("    <main", header_start)
    main_close = "    </main>"
    main_end = page.rindex(main_close)
    footer = page[main_end + len(main_close) :]
    header = page[header_start:main_start]
    if relative_depth == 1:
        header = header.replace("../../", "../")
        footer = footer.replace("../../", "../")
    return header, footer


def add_reference_reader_script(footer: str) -> str:
    script = '    <script src="../../assets/js/reference-reader.js?v=20260531-mobile-reference-reader"></script>\n'
    return footer.replace("  </body>", script + "  </body>")


def toc_html(toc: list[tuple[str, str]]) -> str:
    items = "".join(f'<li><a href="#{esc(anchor)}">{esc(title)}</a></li>' for anchor, title in toc)
    return f'<ol>{items}</ol>'


def clear_docx_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_font(style, *, name: str = "Calibri", size: int = 11, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def ensure_paragraph_style(
    doc: Document,
    name: str,
    *,
    base: str = "Normal",
    size: int = 11,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    before: int = 0,
    after: int = 6,
    line_spacing: float = 1.1,
):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
    set_font(style, size=size, color=color, bold=bold, italic=italic)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    return style


def add_meta_pair(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="WOEK Meta")
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    p.add_run(value)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_ci_docx(intro: list[dict], body: list[dict], target: Path) -> None:
    doc = Document(str(TEMPLATE))
    clear_docx_body(doc)
    props = doc.core_properties
    props.title = TITLE
    props.subject = SUBTITLE
    props.author = "Natalie Weber"
    props.keywords = "Wirkungsökonomie; Wirkungsfinanzpolitik; Staatsschulden; Wirkungshaushalt; Positive Netto-Wirkung; MMT; Public Purpose"
    props.comments = "CI-DOCX aus WÖk-Template; öffentliche Fassung ist ausschließlich die PDF."

    ensure_paragraph_style(doc, "WOEK Eyebrow", size=10, color="3F6F2A", bold=True, before=0, after=6, line_spacing=1.0)
    ensure_paragraph_style(doc, "WOEK Titel", size=22, color="0B2545", bold=True, before=6, after=8, line_spacing=1.05)
    ensure_paragraph_style(doc, "WOEK Untertitel", size=13, color="334155", before=0, after=16, line_spacing=1.12)
    ensure_paragraph_style(doc, "WOEK Meta", size=10, color="334155", before=0, after=2, line_spacing=1.05)
    ensure_paragraph_style(doc, "WOEK Kernsatz", size=12, color="1F3A5F", bold=True, italic=True, before=10, after=10, line_spacing=1.15)
    ensure_paragraph_style(doc, "WOEK Hinweis", size=10, color="555555", italic=True, before=8, after=8, line_spacing=1.1)

    doc.add_paragraph("WIRKUNGSÖKONOMIE", style="WOEK Eyebrow")
    doc.add_paragraph("ARBEITSPAPIER", style="WOEK Eyebrow")
    doc.add_paragraph(TITLE, style="WOEK Titel")
    doc.add_paragraph(SUBTITLE, style="WOEK Untertitel")
    add_meta_pair(doc, "Autorin", "Natalie Weber")
    add_meta_pair(doc, "Status", STATUS)
    add_meta_pair(doc, "Stand", STAND)
    add_meta_pair(doc, "Quelle", "Wirkungsökonomie")
    add_meta_pair(doc, "Öffentliche Fassung", "PDF und Web-Volltext; DOCX nicht öffentlich zum Download")

    for block in intro:
        if block.get("style") in {"Kernsatz", "Callout"}:
            doc.add_paragraph(block["text"], style="WOEK Kernsatz")

    doc.add_paragraph(
        "Dieses Arbeitspapier ist eine konzeptionelle Arbeitsfassung. Es ersetzt keine rechtliche, steuerliche, finanzielle, haushaltspolitische oder wissenschaftliche Fachprüfung.",
        style="WOEK Hinweis",
    )
    add_page_break(doc)

    for block in body:
        if block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=0, cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
            continue

        style = block.get("style", "Normal")
        text = block["text"]
        if style == "Heading1":
            doc.add_paragraph(text, style="Heading 1")
        elif style == "Heading2":
            doc.add_paragraph(text, style="Heading 2")
        elif style == "Heading3":
            doc.add_paragraph(text, style="Heading 3")
        elif style in {"Kernsatz", "Callout"}:
            doc.add_paragraph(text, style="WOEK Kernsatz")
        elif style == "ListBullet":
            doc.add_paragraph(text, style="List Bullet")
        elif style == "ListNumber":
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text, style="Normal")

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="woek-pdf-") as tmp:
        tmp_path = Path(tmp)
        profile = tmp_path / "lo-profile"
        outdir = tmp_path / "out"
        profile.mkdir()
        outdir.mkdir()
        env = os.environ.copy()
        env.update({
            "HOME": str(profile),
            "TMPDIR": "/private/tmp",
            "TEMP": "/private/tmp",
            "TMP": "/private/tmp",
        })
        cmd = [
            "soffice",
            f"-env:UserInstallation=file://{profile}",
            "--invisible",
            "--headless",
            "--norestore",
            "--convert-to",
            "pdf",
            "--outdir",
            str(outdir),
            str(docx_path),
        ]
        result = subprocess.run(cmd, cwd=str(ROOT), env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False)
        generated = outdir / f"{docx_path.stem}.pdf"
        if result.returncode != 0 or not generated.exists() or generated.stat().st_size == 0:
            raise RuntimeError(f"PDF conversion failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
        shutil.copy2(generated, pdf_path)


def render_document_page(intro: list[dict], body_html: str, toc: list[tuple[str, str]], source_hash: str, pdf_hash: str) -> str:
    header, footer = page_shell_prefix_suffix(2)
    footer = add_reference_reader_script(footer)
    quote = next((b["text"] for b in intro if b.get("style") in {"Callout", "Kernsatz"}), "")
    definition = next((b["text"] for b in intro if b.get("style") == "Normal" and b["text"].startswith("Wirkungsfinanzpolitik ist")), "")
    return f"""<!DOCTYPE html>
<html lang="de">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{esc(TITLE)} | Wirkungsökonomie</title>
    <meta name="description" content="{esc(DESCRIPTION)}">
    <meta name="search_title" content="{esc(TITLE)}">
    <meta name="search_description" content="{esc(DESCRIPTION)}">
    <meta name="search_section" content="Dokumente">
    <meta name="search_type" content="Arbeitspapier">
    <meta name="search_tags" content="Wirkungsfinanzpolitik, Staatsschulden, Wirkungshaushalt, Schuldenbremse, MMT, Functional Finance, Public Purpose, öffentliche Finanzen">
    <link rel="stylesheet" href="../../assets/css/style.css?v=20260612-nav-restore">
  </head>
  <body class="reference-ux-page">
{header}    <main class="reference-work reference-reader workpaper-reader" data-pagefind-body>
      <aside class="document-mini-map" data-search-exclude>
        <h2>Inhalt</h2>
        {toc_html(toc)}
      </aside>
      <article class="article-shell">
        <nav class="breadcrumb"><a href="../">Dokumente</a> / Arbeitspapier</nav>
        <p class="hero-kicker">Arbeitspapier · {esc(STATUS)} · Stand {esc(STAND)}</p>
        <h1>{esc(TITLE)}</h1>
        <p class="lead">{esc(SUBTITLE)}</p>
        <div class="document-reader-tools">
          <a class="btn btn-secondary" href="../">Dokumentenbibliothek</a>
          <a class="btn btn-secondary" href="../../bibliothek/wirkungsfinanzpolitik/">Bibliothekseintrag</a>
          <a class="btn btn-secondary" href="../../werkzeuge/wirkungshaushalt/">Wirkungshaushalt</a>
          <button class="btn btn-secondary" type="button" data-print-page>Drucken</button>
        </div>
        <div class="hero-actions">
          <a class="btn btn-primary" href="../../public/downloads/originals/{esc(PUBLIC_PDF_NAME)}">PDF öffnen</a>
          <a class="btn btn-secondary" href="../../wissen/working-papers/">Working Papers</a>
        </div>
        <section class="callout">
          <h2>Kernaussage</h2>
          <p>{esc(quote)}</p>
          <p>{esc(definition)}</p>
        </section>
        <aside class="citation-note" role="note">
          <p class="card-kicker">Schutzlinie</p>
          <h2>Arbeitsfassung, keine Beratung</h2>
          <p>Der Text ist ein konzeptioneller Aufsatz zur Wirkungsökonomie. Er ersetzt keine rechtliche, steuerliche, finanzielle, haushaltspolitische oder wissenschaftliche Fachprüfung.</p>
        </aside>
        <section class="live-reference-notice">
          <h2>Versionshinweis</h2>
          <p>Source-Version: {esc(SOURCE_VERSION)}. Web-Version: {esc(WEB_VERSION)}. Die öffentliche Downloadfassung ist die aus dem WÖk-CI-Template erzeugte PDF. Quellenhash intern: <code>{esc(source_hash)}</code>. PDF-SHA-256: <code>{esc(pdf_hash)}</code>.</p>
        </section>
        <section class="article-body">
          {body_html}
        </section>
      </article>
      <aside class="reference-context-rail" data-search-exclude>
        <section class="toc-card">
          <p class="hero-kicker">Begriffe</p>
          <h2>Anschlussstellen</h2>
          <div class="model-strip">
            <a href="../../begriffe/wirkungshaushalt/">Wirkungshaushalt</a>
            <a href="../../begriffe/schuldenbremse/">Schuldenbremse</a>
            <a href="../../begriffe/mmt/">MMT</a>
            <a href="../../begriffe/public-purpose/">Public Purpose</a>
            <a href="../../begriffe/functional-finance/">Functional Finance</a>
            <a href="../../begriffe/positive-netto-wirkung/">Positive Netto-Wirkung</a>
            <a href="../../begriffe/wirkungsoekonomie/">Wirkungsökonomie</a>
            <a href="../../begriffe/sustainable-value/">Sustainable Value</a>
          </div>
        </section>
        <section class="toc-card">
          <p class="hero-kicker">Einordnung</p>
          <h2>Warum relevant?</h2>
          <p>Der Aufsatz verschiebt die Schuldenfrage von der reinen Betragsperspektive zur Wirkungsbilanz: Entscheidend ist, welche realen Zustände öffentliche Finanzierung verändert.</p>
        </section>
      </aside>
    </main>{footer}
"""


def render_library_page(pdf_hash: str) -> str:
    header, footer = page_shell_prefix_suffix(2)
    return f"""<!DOCTYPE html>
<html lang="de">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{esc(TITLE)} | Bibliothek der Wirkungsökonomie</title>
    <meta name="description" content="{esc(DESCRIPTION)}">
    <meta name="search_title" content="{esc(TITLE)} | Bibliothek der Wirkungsökonomie">
    <meta name="search_description" content="{esc(DESCRIPTION)}">
    <meta name="search_section" content="Bibliothek">
    <meta name="search_type" content="Dokument">
    <link rel="stylesheet" href="../../assets/css/style.css?v=20260612-nav-restore">
  </head>
  <body>
{header}    <main data-pagefind-body>
      <section class="hero compact-hero document-detail-hero">
        <p class="hero-kicker">working-paper · arbeitsfassung</p>
        <h1>{esc(TITLE)}</h1>
        <p class="hero-subtitle">{esc(SUBTITLE)}</p>
        <div class="document-card-badges"><span class="status-badge status-badge--working-paper">working-paper</span><span class="status-badge status-badge--arbeitsfassung">arbeitsfassung</span><span class="status-badge status-badge--expert">Niveau: fortgeschritten</span></div>
      </section>
      <section class="section document-detail-grid">
        <article class="document-detail-main">
          <div class="callout"><strong>Statushinweis:</strong> Dieses Dokument ist eine Arbeitsfassung vom {esc(STAND)} und kann redaktionell weiterentwickelt werden.</div>
          <div class="callout warning"><strong>Schutzlinie:</strong> Keine Rechts-, Steuer-, Finanz-, Anlage- oder Politikberatung.</div>
          <h2>Kurz gesagt</h2>
          <p>Das Arbeitspapier entwickelt Wirkungsfinanzpolitik als Blick auf öffentliche Finanzen: Nicht die bloße Schuldenhöhe entscheidet, sondern die Netto-Wirkung von Einnahmen, Ausgaben, Krediten, Zinsen, Steuern und Investitionen auf Mensch, Planet und Demokratie.</p>
          <h2>Was dich erwartet</h2>
          <p>Eine Brücke zwischen Staatsschuldendebatte, MMT, Functional Finance, Public Purpose, Schuldenbremse, Wirkungshaushalt, planetaren Grenzen und demokratischer Steuerung.</p>
          <h2>Welche Fragen beantwortet das Dokument?</h2>
          <ul><li>Wann sind öffentliche Schulden aus WÖk-Sicht legitim?</li><li>Wie unterscheidet man Wirkschulden, Blindschulden, Verlustschulden und Reparaturschulden?</li><li>Wie sähe ein Wirkungshaushalt als Ergänzung zur Schuldenbremse aus?</li></ul>
          <h2>Für wen geeignet?</h2>
          <p>Politik, Verwaltung, Wissenschaft, Journalismus, öffentliche Haushalte, Stiftungen und alle, die Finanzpolitik nach Zukunftswirkung bewerten wollen.</p>
          <h2>Verwandte Inhalte</h2>
          <ul><li><a href="../../werkzeuge/wirkungshaushalt/">Wirkungshaushalt</a></li><li><a href="../../begriffe/wirkungshaushalt/">Glossar: Wirkungshaushalt</a></li><li><a href="../../wirkungsfelder/staat-recht-demokratie/">Staat, Recht &amp; Demokratie</a></li><li><a href="../../fuer/politik.html">Für Politik &amp; Verwaltung</a></li></ul>
        </article>
        <aside class="document-detail-aside" data-search-exclude>
          <dl>
            <dt>Dokumentart</dt><dd>working-paper</dd>
            <dt>Status</dt><dd>arbeitsfassung</dd>
            <dt>Umfang</dt><dd>Web-Volltext · PDF</dd>
            <dt>Stand / Version</dt><dd>{esc(STAND)} · {esc(WEB_VERSION)}</dd>
            <dt>Zielgruppe</dt><dd>Politik, Verwaltung, Wissenschaft, Journalismus</dd>
            <dt>Niveau</dt><dd>fortgeschritten</dd>
            <dt>PDF-SHA-256</dt><dd><code>{esc(pdf_hash[:16])}...</code></dd>
          </dl>
          <div class="document-chip-row"><span>Wirkungsfinanzpolitik</span><span>Staatsschulden</span><span>Wirkungshaushalt</span><span>Schuldenbremse</span><span>MMT</span><span>Functional Finance</span><span>Public Purpose</span></div>
          <div class="document-action-row"><a class="btn btn-secondary" href="../../dokumente/wirkungsfinanzpolitik/">Online lesen</a><a class="btn btn-primary" href="../../public/downloads/originals/{esc(PUBLIC_PDF_NAME)}">PDF öffnen</a></div>
          <a class="text-link" href="../">Zur Bibliothek</a>
        </aside>
      </section>
    </main>{footer}
"""


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    DOCUMENT_DIR.mkdir(parents=True, exist_ok=True)
    LIBRARY_DIR.mkdir(parents=True, exist_ok=True)
    source_hash = digest(SOURCE)
    intro, body = split_source(read_docx_blocks(SOURCE))
    build_ci_docx(intro, body, INTERNAL_DOCX_TARGET)
    convert_docx_to_pdf(INTERNAL_DOCX_TARGET, PUBLIC_PDF_TARGET)
    pdf_hash = digest(PUBLIC_PDF_TARGET)
    body_html, toc = render_body(body)
    (DOCUMENT_DIR / "index.html").write_text(render_document_page(intro, body_html, toc, source_hash, pdf_hash), encoding="utf-8")
    (LIBRARY_DIR / "index.html").write_text(render_library_page(pdf_hash), encoding="utf-8")
    print(f"Imported {TITLE}: {len(toc)} sections, source={source_hash}, pdf={pdf_hash}")


if __name__ == "__main__":
    main()
