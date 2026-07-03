#!/usr/bin/env python3
"""Export a WÖk study-script Markdown master to a simple DOCX raw draft."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt


TABLE_BORDER = "Table Grid"


def clean_inline(text: str) -> str:
    text = re.sub(r"\$\$([^$]+)\$\$", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = clean_formula_text(text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def clean_formula_text(text: str) -> str:
    replacements = {
        r"\sum": "sum",
        r"\cdot": "*",
        r"\Delta": "Delta ",
        r"\frac": "frac",
        r"\leq": "<=",
        r"\geq": ">=",
        r"\times": "*",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = re.sub(r"\s+", " ", text)
    return text


def set_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 20, 20, 6),
        ("Heading 2", 16, 18, 6),
        ("Heading 3", 14, 16, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, source: Path) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)

    meta = doc.add_paragraph()
    meta.add_run("Rohfassung fuer Claude / CI-CD-Finalisierung").italic = True
    doc.add_paragraph(
        "Quellenbasis: Die neue Ordnung des Wohlstands, WÖk-Referenz, Glossar, WÖk-Werkzeuge, Journal und zitierte externe Fachquellen."
    )


def is_table_start(lines: list[str], idx: int) -> bool:
    return (
        idx + 1 < len(lines)
        and "|" in lines[idx]
        and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[idx + 1])
        is not None
    )


def split_table_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [clean_inline(cell.strip()) for cell in row.split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = TABLE_BORDER

    for idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(width):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].text = row[idx] if idx < len(row) else ""

    doc.add_paragraph()


def export(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    title = next((clean_inline(line.lstrip("# ")) for line in lines if line.startswith("# ")), markdown_path.stem)

    doc = Document()
    set_style(doc)
    add_title(doc, title, markdown_path)

    idx = 0
    in_formula = False
    in_code_block = False
    formula_buffer: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            doc.add_paragraph(clean_inline(" ".join(paragraph_buffer)))
            paragraph_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            flush_paragraph()
            in_code_block = not in_code_block
            idx += 1
            continue

        if in_code_block:
            idx += 1
            continue

        if line.startswith("# "):
            flush_paragraph()
            idx += 1
            continue

        if line.strip() == "$$":
            flush_paragraph()
            if in_formula:
                doc.add_paragraph("Formel: " + clean_formula_text(" ".join(formula_buffer)), style="Intense Quote")
                formula_buffer = []
                in_formula = False
            else:
                in_formula = True
            idx += 1
            continue

        if in_formula:
            formula_buffer.append(line.strip())
            idx += 1
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            flush_paragraph()
            alt, path = image_match.groups()
            doc.add_paragraph(f"Bildvorgabe: {clean_inline(alt)} — {path}", style="Intense Quote")
            idx += 1
            continue

        if is_table_start(lines, idx):
            flush_paragraph()
            table_rows = [split_table_row(line)]
            idx += 2
            while idx < len(lines) and "|" in lines[idx] and lines[idx].strip():
                table_rows.append(split_table_row(lines[idx]))
                idx += 1
            add_markdown_table(doc, table_rows)
            continue

        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            idx += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(clean_inline(stripped[4:]), style="Heading 3")
        elif stripped.startswith("## "):
            flush_paragraph()
            doc.add_paragraph(clean_inline(stripped[3:]), style="Heading 2")
        elif stripped.startswith("# "):
            flush_paragraph()
            doc.add_paragraph(clean_inline(stripped[2:]), style="Heading 1")
        elif re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            doc.add_paragraph(clean_inline(re.sub(r"^[-*]\s+", "", stripped)), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            doc.add_paragraph(clean_inline(stripped))
        else:
            paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    export(args.markdown, args.out)


if __name__ == "__main__":
    main()
