#!/usr/bin/env python3
import html
import json
import re
import shutil
import textwrap
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

ROOT = Path(__file__).resolve().parents[2]
SOURCE_ROOT = Path("/Users/hagen/Documents/Rechner-Cleanup/Sortiert_2026-05-25/02_Dokumente")
REGISTRY_PATH = ROOT / "assets/data/document-registry.json"
IMPORTS_PATH = ROOT / "assets/data/local-document-imports.json"
CONTENT_DIR = ROOT / "assets/data/document-online"
PDF_DIR = ROOT / "assets/pdf/imported"
AUDIT_PATH = ROOT / "docs/local-document-import-audit.md"

PRIVATE_PATTERNS = re.compile(
    r"Finanzen-Amtliches|Mietvertrag|Kaufvertrag|Rechnung|HUK24|Bankauszug|Stammbuch|"
    r"Aufhebungsvertrag|Firmenwagen|Application Fisbeck|Projektmanagementplan|Adesso|EnBW|"
    r"SAP|PM_TL_|PITPM|Business_Requirements|Functional_Specification|Seminararbeit|Bewerbung-CV",
    re.I,
)
RELEVANCE_PATTERNS = re.compile(
    r"wirkung|woek|wirkungs|nachhalt|marketing|t-sroi|sdg|esg|steuer|dossier|working|"
    r"whitepaper|konzept|impact|resilienz|demokratie|medien|klima|energie|rente|arbeit|"
    r"produkt|wohnen|zange|apfel|milram|ki|transformation|integritaet|desinformation|lieferkette",
    re.I,
)


def norm_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def extract_docx_paragraphs(path: Path):
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = norm_text(para.text)
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [norm_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                paragraphs.append(row_text)
    return paragraphs


def extract_pdf_paragraphs(path: Path):
    reader = PdfReader(str(path))
    paragraphs = []
    for page in reader.pages[:80]:
        text = page.extract_text() or ""
        for block in re.split(r"\n{2,}", text):
            clean = norm_text(block)
            if clean:
                paragraphs.append(clean)
    return paragraphs


def paragraphs_for(path: Path):
    if path.suffix.lower() == ".docx":
        return extract_docx_paragraphs(path)
    if path.suffix.lower() == ".pdf":
        return extract_pdf_paragraphs(path)
    if path.suffix.lower() in {".md", ".html"}:
        return [norm_text(path.read_text(encoding="utf-8", errors="ignore"))]
    return []


def html_fragment(title: str, paragraphs):
    chunks = []
    for idx, paragraph in enumerate(paragraphs):
      if not paragraph:
          continue
      escaped = html.escape(paragraph)
      if idx == 0 and len(paragraph) < 140:
          chunks.append(f"<h3>{escaped}</h3>")
      elif len(paragraph) < 95 and not paragraph.endswith("."):
          chunks.append(f"<h3>{escaped}</h3>")
      elif " | " in paragraph and len(paragraph) < 800:
          cells = [html.escape(cell.strip()) for cell in paragraph.split(" | ") if cell.strip()]
          chunks.append("<ul>" + "".join(f"<li>{cell}</li>" for cell in cells) + "</ul>")
      else:
          chunks.append(f"<p>{escaped}</p>")
    if not chunks:
        chunks.append(f"<p>{html.escape(title)} wird als Onlinefassung bereitgestellt. Die PDF-Fassung steht ergänzend zum Download bereit.</p>")
    return "\n".join(chunks)


def make_pdf_from_paragraphs(out_path: Path, title: str, paragraphs):
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, leftMargin=56, rightMargin=56, topMargin=56, bottomMargin=56)
    styles = getSampleStyleSheet()
    story = [Paragraph(html.escape(title), styles["Title"]), Spacer(1, 12)]
    for paragraph in paragraphs:
        if not paragraph:
            continue
        if len(paragraph) < 95 and not paragraph.endswith("."):
            story.extend([Spacer(1, 8), Paragraph(html.escape(paragraph), styles["Heading2"])])
        else:
            wrapped = html.escape(paragraph)
            story.append(Paragraph(wrapped, styles["BodyText"]))
            story.append(Spacer(1, 7))
    doc.build(story)


def file_size(path: Path) -> str:
    size = path.stat().st_size
    if size >= 1024 * 1024:
        return f"{size / (1024 * 1024):.1f} MB"
    return f"{size / 1024:.0f} KB"


def scan_candidates():
    files = []
    if not SOURCE_ROOT.exists():
        return files
    for path in SOURCE_ROOT.rglob("*"):
        if path.suffix.lower() not in {".docx", ".pdf", ".md", ".html"}:
            continue
        text = str(path)
        if PRIVATE_PATTERNS.search(text):
            decision = "ausgeschlossen"
            reason = "private, amtliche oder fremde Arbeitsdatei nach Pfad/Dateiname"
        elif RELEVANCE_PATTERNS.search(text):
            decision = "review"
            reason = "WÖk-/ESG-/SDG-/Konzept-Bezug nach Dateiname"
        else:
            decision = "ausgeschlossen"
            reason = "kein klarer WÖk-Bezug im Dateinamen"
        files.append({"path": str(path), "decision": decision, "reason": reason})
    return files


def main():
    CONTENT_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)

    registry = json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
    imports = json.loads(IMPORTS_PATH.read_text(encoding="utf-8"))
    by_id = {item["id"]: item for item in registry}
    imported = []
    blocked = []

    for item in imports:
        source = Path(item["sourcePath"])
        if not source.exists():
            blocked.append((item, "Quelle fehlt"))
            continue
        paragraphs = paragraphs_for(source)
        content_path = CONTENT_DIR / f"{item['slug']}.html"
        content_path.write_text(html_fragment(item["title"], paragraphs), encoding="utf-8")

        out_pdf = PDF_DIR / f"{item['slug']}.pdf"
        preferred_value = item.get("preferredPdfPath")
        preferred = Path(preferred_value) if preferred_value else None
        if preferred and preferred.exists() and preferred.is_file():
            shutil.copy2(preferred, out_pdf)
            pdf_note = "bestehende PDF übernommen"
        else:
            make_pdf_from_paragraphs(out_pdf, item["title"], paragraphs)
            pdf_note = "PDF aus Textauszug erzeugt"

        entry = {
            "id": item["id"],
            "slug": item["slug"],
            "title": item["title"],
            "type": item["type"],
            "category": item["category"],
            "status": "current",
            "stand": item["stand"],
            "summary": item["summary"],
            "audience": item["audience"],
            "keyPoints": [
                item["summary"],
                "Die Onlinefassung macht das Dokument direkt lesbar.",
                "Die PDF-Fassung steht ergänzend als Download bereit."
            ],
            "onlineUrl": f"/bibliothek/{item['slug']}/",
            "sourceOnlineUrl": item.get("sourceOnlineUrl"),
            "pdfUrl": f"/assets/pdf/imported/{item['slug']}.pdf",
            "docxUrl": None,
            "fileSize": file_size(out_pdf),
            "contentHtmlPath": f"assets/data/document-online/{item['slug']}.html",
            "relatedTerms": item.get("relatedTerms", []),
            "relatedFields": item.get("relatedFields", []),
            "relatedTools": item.get("relatedTools", []),
            "relatedPages": item.get("relatedPages", []),
            "isArchive": False,
            "isPublic": True,
            "importSource": str(source),
        }
        by_id[item["id"]] = entry
        imported.append((entry, pdf_note, len(paragraphs)))

    merged = list(by_id.values())
    merged.sort(key=lambda item: (item.get("isArchive", False), item.get("category", [""])[0], item["title"]))
    REGISTRY_PATH.write_text(json.dumps(merged, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    candidates = scan_candidates()
    review = [item for item in candidates if item["decision"] == "review"]
    imported_paths = {entry["importSource"] for entry, _, _ in imported}
    for item in imports:
        if item.get("preferredPdfPath"):
            imported_paths.add(item["preferredPdfPath"])
    still_review = [item for item in review if item["path"] not in imported_paths]

    audit_lines = [
        "# Local Document Import Audit",
        "",
        "Stand: 2026-05-26",
        "",
        "## Zusammenfassung",
        "",
        f"- Dateien im lokalen Dokumentenbestand: {len(candidates)}",
        f"- Review-Kandidaten nach Heuristik: {len(review)}",
        f"- In diesem Lauf importiert: {len(imported)}",
        f"- Blockiert: {len(blocked)}",
        f"- Weiter review-pflichtig: {len(still_review)}",
        "",
        "Die Heuristik schließt private, amtliche und offensichtlich fremde Dateien aus. Importiert werden nur kuratierte WÖk-nahe Dokumente aus der Allowlist in `assets/data/local-document-imports.json`.",
        "",
        "## Importiert",
        "",
        "| id | Titel | PDF | Absätze | Quelle |",
        "| --- | --- | --- | ---: | --- |",
    ]
    for entry, pdf_note, para_count in imported:
        audit_lines.append(f"| {entry['id']} | {entry['title']} | {pdf_note} | {para_count} | `{entry['importSource']}` |")

    audit_lines.extend(["", "## Blockiert", ""])
    if blocked:
        for item, reason in blocked:
            audit_lines.append(f"- `{item['sourcePath']}` - {reason}")
    else:
        audit_lines.append("- Keine")

    audit_lines.extend(["", "## Weiter review-pflichtige Kandidaten", ""])
    for item in still_review[:120]:
        audit_lines.append(f"- `{item['path']}` - {item['reason']}")
    if len(still_review) > 120:
        audit_lines.append(f"- ... {len(still_review) - 120} weitere Kandidaten")

    AUDIT_PATH.write_text("\n".join(audit_lines) + "\n", encoding="utf-8")
    print(f"Imported {len(imported)} documents; {len(still_review)} candidates remain for review.")


if __name__ == "__main__":
    main()
