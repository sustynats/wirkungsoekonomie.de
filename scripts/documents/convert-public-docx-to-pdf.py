#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path.cwd()
AUDIT = ROOT / "docs/public-docx-conversion-audit.md"
PUBLIC_ROOTS = [ROOT / "assets/downloads", ROOT / "public/downloads"]


def normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def docx_text(path: Path) -> list[str]:
    document = Document(path)
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return paragraphs


def pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for paragraph in paragraphs:
        story.append(Paragraph(paragraph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
        story.append(Spacer(1, 6))
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=42, leftMargin=42, topMargin=48, bottomMargin=48).build(story)


def main() -> None:
    docx_files = sorted(
        file
        for root in PUBLIC_ROOTS
        if root.exists()
        for file in root.rglob("*")
        if file.suffix.lower() in {".docx", ".doc"}
    )
    generated = []
    existing = []
    failed = []
    passed_compare = []
    review = []

    for source in docx_files:
        target = source.with_suffix(".pdf")
        if target.exists():
            existing.append((source, target))
            continue
        try:
            paragraphs = docx_text(source)
            write_pdf(target, source.stem.replace("_", " "), paragraphs)
            generated.append((source, target))
            try:
                original = normalize("\n".join(paragraphs))
                converted = normalize(pdf_text(target))
                if original and original[:1000] in converted:
                    passed_compare.append((source, target))
                else:
                    review.append((source, target, "Textvergleich wegen PDF-Zeilenumbrüchen review-pflichtig"))
            except Exception as exc:  # noqa: BLE001
                review.append((source, target, f"Textvergleich nicht stabil: {exc}"))
        except Exception as exc:  # noqa: BLE001
            failed.append((source, target, str(exc)))

    lines = [
        "# Public DOCX to PDF Conversion Audit",
        "",
        "## Zusammenfassung",
        "",
        f"- Öffentliche DOCX-/Word-Quellen gefunden: {len(docx_files)}",
        f"- PDF bereits vorhanden: {len(existing)}",
        f"- PDF neu erzeugt: {len(generated)}",
        f"- Textvergleich bestanden: {len(passed_compare)}",
        f"- Review-pflichtig: {len(review)}",
        f"- Fehlgeschlagen: {len(failed)}",
        "",
        "Die Konvertierung verändert keine Quelldokumente. DOCX-Dateien bleiben nur bis zur anschließenden Entfernung aus öffentlichen Asset-Pfaden als interne Quellen im Arbeitsbaum vorhanden.",
        "",
        "## Neu erzeugte PDFs",
        "",
    ]
    if generated:
        lines.extend(f"- `{src.relative_to(ROOT)}` -> `{dst.relative_to(ROOT)}`" for src, dst in generated)
    else:
        lines.append("- Keine")
    lines.extend(["", "## Bereits vorhandene PDFs", ""])
    if existing:
        lines.extend(f"- `{src.relative_to(ROOT)}` -> `{dst.relative_to(ROOT)}`" for src, dst in existing)
    else:
        lines.append("- Keine")
    lines.extend(["", "## Review-pflichtig", ""])
    if review:
        lines.extend(f"- `{src.relative_to(ROOT)}` -> `{dst.relative_to(ROOT)}`: {reason}" for src, dst, reason in review)
    else:
        lines.append("- Keine")
    lines.extend(["", "## Fehlgeschlagen", ""])
    if failed:
        lines.extend(f"- `{src.relative_to(ROOT)}`: {reason}" for src, _dst, reason in failed)
    else:
        lines.append("- Keine")
    AUDIT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"Public DOCX conversion: {len(generated)} generated, {len(existing)} existing, {len(review)} review, {len(failed)} failed -> docs/public-docx-conversion-audit.md")
    if failed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
