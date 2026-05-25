#!/usr/bin/env python3
"""Small Markdown to DOCX fallback for WÖk dossier layout runs.

Pandoc remains the preferred converter. This fallback exists so the layout
standardization pipeline can still produce auditable DOCX files when Pandoc is
not installed. It intentionally supports the Markdown subset used by layout
test and dossier source files: frontmatter, headings, paragraphs, blockquotes,
lists, simple pipe tables and fenced code blocks.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND = {
    "navy": RGBColor(0x0B, 0x10, 0x20),
    "ivory": RGBColor(0xF6, 0xF1, 0xE8),
    "green": RGBColor(0x2F, 0x7D, 0x5C),
    "gold": RGBColor(0xC8, 0x9B, 0x3C),
    "coral": RGBColor(0xC8, 0x5A, 0x4A),
    "text": RGBColor(0x22, 0x24, 0x2C),
}

REQUIRED_STYLES = [
    "Title",
    "Subtitle",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Normal",
    "Quote",
    "List Paragraph",
    "Caption",
    "Table Text",
    "Footer",
    "Header",
]


def parse_frontmatter(text: str) -> tuple[dict[str, str | bool], str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end < 0:
        return {}, text
    block = text[3:end].strip()
    body = text[end + 4 :].lstrip("\r\n")
    meta: dict[str, str | bool] = {}
    for line in block.splitlines():
        match = re.match(r"^([A-Za-z0-9_-]+):\s*(.*)$", line)
        if not match:
            continue
        key, value = match.group(1), match.group(2).strip().strip("\"'")
        if value.lower() in {"true", "false"}:
            meta[key] = value.lower() == "true"
        else:
            meta[key] = value
    return meta, body


def ensure_style(document: Document, name: str, style_type: WD_STYLE_TYPE = WD_STYLE_TYPE.PARAGRAPH):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, style_type)


def set_run_font(run, name: str, size: int | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Merriweather"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Merriweather")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BRAND["text"]

    for style_name, size in [("Title", 26), ("Subtitle", 14), ("Heading 1", 19), ("Heading 2", 15), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Playfair Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Playfair Display")
        style.font.size = Pt(size)
        style.font.color.rgb = BRAND["navy"]

    quote = document.styles["Quote"]
    quote.font.name = "Merriweather"
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Merriweather")
    quote.font.italic = True
    quote.font.color.rgb = BRAND["green"]

    table_text = ensure_style(document, "Table Text")
    table_text.font.name = "Source Sans 3"
    table_text._element.rPr.rFonts.set(qn("w:eastAsia"), "Source Sans 3")
    table_text.font.size = Pt(9)

    caption = ensure_style(document, "Caption")
    caption.font.name = "Source Sans 3"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Source Sans 3")
    caption.font.size = Pt(9)
    caption.font.color.rgb = BRAND["green"]


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title_page(document: Document, meta: dict[str, str | bool], signet: Path | None) -> None:
    if not meta.get("title") and not meta.get("subtitle"):
        return
    if signet and signet.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        try:
            run.add_picture(str(signet), width=Inches(0.72))
        except Exception:
            pass

    title = str(meta.get("title") or "")
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(title)

    subtitle = str(meta.get("subtitle") or "")
    if subtitle:
        document.add_paragraph(subtitle, style="Subtitle")

    metadata = [
        ("Autorin", str(meta.get("author") or meta.get("autorin") or "Natalie Weber")),
        ("Referenz", str(meta.get("reference") or "Wirkungsökonomie")),
        ("Version", str(meta.get("version") or "")),
        ("Stand", str(meta.get("stand") or "")),
        ("Status", str(meta.get("status") or "")),
    ]
    for label, value in metadata:
        if not value:
            continue
        p = document.add_paragraph()
        p.style = "Normal"
        r = p.add_run(f"{label}: ")
        set_run_font(r, "Source Sans 3", 9, BRAND["green"], True)
        r2 = p.add_run(value)
        set_run_font(r2, "Source Sans 3", 9, BRAND["text"])

    document.add_section(WD_SECTION_START.NEW_PAGE)


def clear_document_body(document: Document) -> None:
    body = document._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
            continue
        body.remove(child)
    if sect_pr is not None and sect_pr.getparent() is None:
        body.append(sect_pr)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_markdown_body(document: Document, body: str) -> None:
    lines = body.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    in_fence = False
    code_lines: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line.startswith("```"):
            if in_fence:
                if code_lines:
                    document.add_paragraph("\n".join(code_lines), style="Caption")
                code_lines = []
                in_fence = False
            else:
                in_fence = True
            i += 1
            continue

        if in_fence:
            code_lines.append(raw.rstrip())
            i += 1
            continue

        if not line:
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            document.add_paragraph(clean_inline(heading.group(2).strip()), style=f"Heading {level}")
            i += 1
            continue

        if line.startswith("|"):
            separator_index = i + 1
            while separator_index < len(lines) and not lines[separator_index].strip():
                separator_index += 1
        if line.startswith("|") and separator_index < len(lines) and is_table_separator(lines[separator_index]):
            header = split_table_row(line)
            rows: list[list[str]] = []
            i = separator_index + 1
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line:
                    i += 1
                    continue
                if not row_line.startswith("|"):
                    break
                if not is_table_separator(row_line):
                    row_values = split_table_row(row_line)
                    if len(row_values) != len(header):
                        break
                    rows.append(row_values)
                i += 1
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for idx, value in enumerate(header):
                cell = table.rows[0].cells[idx]
                shade_cell(cell, "F6F1E8")
                cell.text = value
                for p in cell.paragraphs:
                    p.style = "Table Text"
                    for run in p.runs:
                        set_run_font(run, "Source Sans 3", 9, BRAND["navy"], True)
            for row_values in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row_values[: len(cells)]):
                    cells[idx].text = value
                    for p in cells[idx].paragraphs:
                        p.style = "Table Text"
            continue

        if line.startswith(">"):
            document.add_paragraph(clean_inline(re.sub(r"^>\s?", "", line)), style="Quote")
            i += 1
            continue

        if re.match(r"^\*\*[^*]+:\*\*", line):
            document.add_paragraph(clean_inline(line), style="Normal")
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        ordered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if bullet or ordered:
            text = (bullet or ordered).group(1)
            document.add_paragraph(clean_inline(text), style="List Bullet" if bullet else "List Number")
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith(">")
                or next_line.startswith("|")
                or next_line.startswith("```")
                or re.match(r"^[-*+]\s+", next_line)
                or re.match(r"^\d+[.)]\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        document.add_paragraph(clean_inline(" ".join(paragraph_lines)), style="Normal")


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def inspect_reference(path: Path) -> int:
    document = Document(path)
    styles = {style.name for style in document.styles}
    result = {
        "path": str(path),
        "fileType": path.suffix.lower().lstrip("."),
        "usableAsReferenceDoc": path.suffix.lower() == ".docx",
        "styles": {name: name in styles for name in REQUIRED_STYLES},
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


def convert_docx(input_path: Path, output_path: Path) -> int:
    document = Document(str(input_path))
    configure_styles(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return 0


def convert(input_path: Path, output_path: Path, reference_doc: Path | None, signet: Path | None) -> int:
    if input_path.suffix.lower() == ".docx":
        return convert_docx(input_path, output_path)
    text = input_path.read_text(encoding="utf-8", errors="replace")
    meta, body = parse_frontmatter(text)
    if not meta.get("title"):
        first_heading = re.search(r"(?m)^#\s+(.+)$", body)
        if first_heading and not body[: first_heading.start()].strip():
            meta["title"] = first_heading.group(1).strip()
            body = body[: first_heading.start()] + body[first_heading.end() :]
    document = Document(str(reference_doc)) if reference_doc and reference_doc.exists() else Document()
    clear_document_body(document)
    configure_styles(document)
    add_title_page(document, meta, signet)
    add_markdown_body(document, body)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return 0


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", nargs="?")
    parser.add_argument("output", nargs="?")
    parser.add_argument("--reference-doc", default="")
    parser.add_argument("--signet", default="")
    parser.add_argument("--inspect-reference", action="store_true")
    args = parser.parse_args()

    reference = Path(args.reference_doc) if args.reference_doc else None
    if args.inspect_reference:
        if not reference:
            print("missing --reference-doc", file=sys.stderr)
            return 2
        return inspect_reference(reference)
    if not args.input or not args.output:
        print("input and output are required", file=sys.stderr)
        return 2
    return convert(
        Path(args.input),
        Path(args.output),
        reference,
        Path(args.signet) if args.signet else None,
    )


if __name__ == "__main__":
    sys.exit(main())
