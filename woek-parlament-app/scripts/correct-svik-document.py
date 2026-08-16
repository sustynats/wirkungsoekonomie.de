#!/usr/bin/env python3
"""Apply the approved 1.1 wording corrections to the SVIK publication."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


AUTHOR = "Institut für Wirkungsökonomie"
TITLE = "Wirkungsökonomische Analyse des Sondervermögens Infrastruktur und Klimaneutralität"
SUBJECT = "Wirkungspotenziale, Wirkungsrisiken, Schutzgrenzen und Rückkopplung"


def set_metadata(document: Document) -> None:
    properties = document.core_properties
    properties.author = AUTHOR
    properties.last_modified_by = AUTHOR
    properties.title = TITLE
    properties.subject = SUBJECT
    properties.keywords = "Wirkungsökonomie; Sondervermögen; Infrastruktur; Klimaneutralität; Wirkungspotenzial"
    properties.comments = ""
    properties.category = "Fachanalyse"


def set_publication_footers(document: Document) -> None:
    """Name both author and institutional publisher without destroying page fields."""
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Institut für Wirkungsökonomie" in paragraph.text:
                continue
            replace_in_run(
                paragraph,
                "Natalie Weber  •",
                "Natalie Weber  •  Institut für Wirkungsökonomie  •",
            )


def mark_real_table_headers(document: Document) -> None:
    """Give every Word table a detectable header row for assistive technology."""
    for table in document.tables:
        if not table.rows:
            continue
        tr_properties = table.rows[0]._tr.get_or_add_trPr()
        if tr_properties.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader") is None:
            tr_properties.append(OxmlElement("w:tblHeader"))


def replace_in_run(paragraph, old: str, new: str) -> bool:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    return False


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--metadata-only", action="store_true")
    args = parser.parse_args()

    document = Document(args.input)
    set_metadata(document)
    mark_real_table_headers(document)
    set_publication_footers(document)

    if not args.metadata_only:
        replacements = []

        replacements.append(replace_in_run(
            document.paragraphs[2],
            "Arbeitsstand zur fachlichen Veröffentlichung",
            "Fachanalyse · Stand 15. August 2026"
        ))

        replacements.append(replace_in_run(
            document.paragraphs[14],
            "Einzelne Programme sind klar positiv",
            "Einzelne Programme zeigen klare positive Wirkungspotenziale"
        ))
        replacements.append(replace_in_run(
            document.paragraphs[133],
            "6. Portfolio-Gesamtbewertung",
            "6. Vorläufige modellierte Ex-ante-WÖk-Matrix"
        ))
        replacements.append(replace_in_run(
            document.paragraphs[258],
            "positive Netto-Wirkung maximieren",
            "das Portfolio auf positive Netto-Wirkung ausrichten"
        ))
        replacements.append(replace_in_run(
            document.paragraphs[361],
            "Führender Begriffsleitfaden der Wirkungsökonomie, Version 1.0.",
            "Führender Begriffsleitfaden der Wirkungsökonomie, Version 1.5."
        ))
        replacements.append(replace_in_run(
            document.paragraphs[366],
            "Whitepaper T-SROI – Transformational Social Return on Investment.",
            "T-SROI-Rechenstandard, Version 1.1."
        ))
        replacements.append(replace_in_run(
            document.paragraphs[369],
            "WÖk-Master-Items, Version 1.2",
            "WÖk-Master-Items, Version 1.3"
        ))

        portfolio_judgement = document.tables[4].rows[0].cells[-1].paragraphs[1]
        replacements.append(replace_in_run(
            portfolio_judgement,
            "Hohes Wirkungspotenzial – bislang nur bedingt nachgewiesene positive Netto-Wirkung.",
            "Hohes modelliertes Wirkungspotenzial – positive Netto-Wirkung bislang nicht hinreichend nachgewiesen."
        ))

        release_gate = document.tables[6].rows[0].cells[-1].paragraphs[1]
        release_gate.runs[1].text = (
            "eine belastbare positive Ex-ante-Netto-Wirkungsbewertung, ein plausibler positiver Wirkpfad, "
            "Schutzgrenzen, Zusätzlichkeit und Umsetzungskapazität"
        )
        release_gate.runs[2].text = (
            " vorliegen; Fortführung und Skalierung müssen von beobachteten Zustandsveränderungen, "
            "Evidenz und nachvollziehbarer Zurechnung abhängen."
        )

        results_chain_row = document.tables[12].rows[6]
        results_chain_row.cells[1].paragraphs[0].runs[0].text = "Wirkung"
        results_chain_row.cells[2].paragraphs[0].runs[0].text = (
            "Tatsächliche Zustandsveränderung für Mensch, Planet oder Demokratie"
        )
        terminology_note = results_chain_row.cells[2].add_paragraph(
            "Outcome und Impact sind hier Begriffe des amtlichen Results-Chain-Modells. "
            "WÖk-Wirkung bezeichnet davon getrennt die tatsächliche Zustandsveränderung; "
            "Transformationswirkung, Nichtkompensation und Rückkopplung sind zusätzliche Prüfebenen."
        )
        terminology_note.style = document.styles["Normal"]

        document.tables[36].rows[4].cells[1].paragraphs[0].runs[0].text = (
            "Auslöser → Wirkungspotenzial/Wirkungsrisiko → Wirkmechanismus und Bedingungen → "
            "Umsetzung/Output → beobachtete Zustandsveränderung → Evidenz und Zurechnung → "
            "Bewertung → Schutz- und Systemprüfung → Transformationswirkung → Rückkopplung und Lernen"
        )

        matrix_heading = document.paragraphs[133]
        matrix_note = matrix_heading.insert_paragraph_before(
            "Keine gemessene Wirkung: Die Matrix zeigt eine modellierte Richtung und Umsetzungsbereitschaft "
            "auf Basis der aktuellen Architektur und Evidenz. Evidenzgrad und Schutzgrenzen bleiben getrennt."
        )
        matrix_heading._p.addnext(matrix_note._p)
        matrix_note.style = document.styles["Normal"]

        if not all(replacements):
            raise RuntimeError("Mindestens eine erwartete Fachformulierung wurde im Ausgangsdokument nicht gefunden.")

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    main()
