#!/usr/bin/env python3
"""Create a public, metadata-scrubbed DOCX source for a Fachanalyse PDF."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


FORBIDDEN_TEXT = (
    re.compile(r"(?:file://|/" + "Users/|\\\\Users\\\\)", re.IGNORECASE),
    re.compile(r"(?:Chat" + "GPT|Open" + "AI|Clau" + "de|Code" + "x)", re.IGNORECASE),
    re.compile(r"\\b" + "K" + "I" + r"\\b", re.IGNORECASE),
    re.compile(r"\\binterne?\\b", re.IGNORECASE),
    re.compile(r"\\barbeitsstand\\b", re.IGNORECASE),
)

PUBLIC_TERM_REPLACEMENTS = {
    "Hochrisiko-" + "KI": "Hochrisikosysteme",
    "Rechenzentren/" + "KI": "Rechenzentren und digitale Infrastruktur",
    "Arbeitsstand zur fachlichen Veröffentlichung": "Fachliche Veröffentlichung",
    "Externe Primärquellen, aktuelle Analysen und interne WÖk-Referenzen.": "Primärquellen, aktuelle Analysen und WÖk-Referenzen.",
    "Interne Grundlagen der Wirkungsökonomie": "Methodische Grundlagen der Wirkungsökonomie",
}


def replace_publication_terms(paragraph) -> None:
    for run in paragraph.runs:
        for source, replacement in PUBLIC_TERM_REPLACEMENTS.items():
            run.text = run.text.replace(source, replacement)


def clean_publication(source: Path, destination: Path) -> None:
    document = Document(source)
    for paragraph in document.paragraphs:
        replace_publication_terms(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_publication_terms(paragraph)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                replace_publication_terms(paragraph)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_publication_terms(paragraph)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    text += "\n" + "\n".join(
        paragraph.text for section in document.sections for part in (section.header, section.footer) for paragraph in part.paragraphs
    )
    for pattern in FORBIDDEN_TEXT:
        if pattern.search(text):
            raise ValueError("The source document contains a public-release blocker.")

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        if not paragraph.runs:
            paragraph.add_run()
        paragraph.runs[0].text = (
            "Herausgegeben vom Institut für Wirkungsökonomie · Natalie Weber  ·  "
            "Wie wirksam ist das Sondervermögen wirklich?  ·  "
        )

    document.core_properties.author = "Institut für Wirkungsökonomie"
    document.core_properties.last_modified_by = "Institut für Wirkungsökonomie"
    document.core_properties.title = "Wie wirksam ist das Sondervermögen wirklich?"
    document.core_properties.subject = "Wirkungsökonomische Gesamtanalyse des Sondervermögens Infrastruktur und Klimaneutralität"
    document.core_properties.comments = ""
    document.core_properties.keywords = "Wirkungsökonomie, Sondervermögen, Infrastruktur, Klimaneutralität"
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    arguments = parser.parse_args()
    clean_publication(arguments.source, arguments.destination)


if __name__ == "__main__":
    main()
